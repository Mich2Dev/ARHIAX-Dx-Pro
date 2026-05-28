"""Binary export helpers for DX Pro executive reports."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


class ReportExportService:
    def __init__(self, root: Path) -> None:
        self.root = root
        self.root.mkdir(parents=True, exist_ok=True)

    def export(
        self,
        case_id: str,
        report_pack: dict[str, Any],
        render_pack: dict[str, Any],
        targets: list[str],
    ) -> list[dict[str, Any]]:
        case_root = self.root / case_id
        case_root.mkdir(parents=True, exist_ok=True)
        exports = []
        markdown = str(render_pack.get("markdown", ""))

        if "markdown" in targets:
            path = case_root / "executive-report.md"
            path.write_text(markdown, encoding="utf-8")
            exports.append(self._file_record("markdown", path))
        if "docx" in targets:
            path = case_root / "executive-report.docx"
            self._write_docx(path, report_pack)
            exports.append(self._file_record("docx", path))
        if "pdf" in targets:
            path = case_root / "executive-report.pdf"
            self._write_pdf(path, report_pack)
            exports.append(self._file_record("pdf", path))
        return exports

    def _write_docx(self, path: Path, report_pack: dict[str, Any]) -> None:
        doc = Document()
        doc.core_properties.title = str(report_pack.get("report_title", "ARHIAX Dx Pro Executive Diagnostic Report"))
        doc.add_heading(str(report_pack.get("report_title", "ARHIAX Dx Pro Executive Diagnostic Report")), level=0)
        client = report_pack.get("client") or {}
        doc.add_paragraph(f"Cliente: {client.get('name', 'Client')}")
        doc.add_paragraph(f"Engagement: {report_pack.get('engagement_id', 'unknown')}")
        doc.add_paragraph(f"Estado: {report_pack.get('report_status', 'consultant_review_required')}")
        doc.add_heading("Tesis Ejecutiva", level=1)
        doc.add_paragraph(str(report_pack.get("executive_thesis", "")))
        for section in report_pack.get("sections") or []:
            doc.add_heading(str(section.get("title", "Section")), level=1)
            doc.add_paragraph(str(section.get("body", "")))
        if report_pack.get("exhibits"):
            doc.add_heading("Exhibits", level=1)
            for exhibit in report_pack.get("exhibits") or []:
                doc.add_heading(str(exhibit.get("title", exhibit.get("id", "Exhibit"))), level=2)
                doc.add_paragraph(str(exhibit.get("type", "data")))
                doc.add_paragraph(str(exhibit.get("data", {})))
        if report_pack.get("appendices"):
            doc.add_heading("Anexos", level=1)
            for appendix in report_pack.get("appendices") or []:
                doc.add_heading(str(appendix.get("title", appendix.get("id", "Appendix"))), level=2)
                doc.add_paragraph(str(appendix.get("content", {})))
        doc.save(path)

    def _write_pdf(self, path: Path, report_pack: dict[str, Any]) -> None:
        font_name = self._register_pdf_font()
        pdf = canvas.Canvas(str(path), pagesize=LETTER)
        width, height = LETTER
        y = height - inch
        lines = self._pdf_lines(report_pack)
        pdf.setTitle(str(report_pack.get("report_title", "ARHIAX Dx Pro Executive Diagnostic Report")))
        pdf.setFont(font_name, 16)
        pdf.drawString(inch, y, str(report_pack.get("report_title", "ARHIAX Dx Pro Executive Diagnostic Report")))
        y -= 0.4 * inch
        pdf.setFont(font_name, 10)
        for line in lines:
            if y < inch:
                pdf.showPage()
                pdf.setFont(font_name, 10)
                y = height - inch
            pdf.drawString(inch, y, line[:110])
            y -= 0.22 * inch
        pdf.save()

    def _register_pdf_font(self) -> str:
        try:
            candidate = Path("C:/Windows/Fonts/arial.ttf")
            if candidate.exists():
                pdfmetrics.registerFont(TTFont("ArialUnicode", str(candidate)))
                return "ArialUnicode"
        except Exception:
            pass
        return "Helvetica"

    def _pdf_lines(self, report_pack: dict[str, Any]) -> list[str]:
        client = report_pack.get("client") or {}
        lines = [
            f"Cliente: {client.get('name', 'Client')}",
            f"Engagement: {report_pack.get('engagement_id', 'unknown')}",
            f"Estado: {report_pack.get('report_status', 'consultant_review_required')}",
            "",
            "Tesis Ejecutiva",
            str(report_pack.get("executive_thesis", "")),
            "",
        ]
        for section in report_pack.get("sections") or []:
            lines.append(str(section.get("title", "Section")))
            lines.extend(self._wrap(str(section.get("body", ""))))
            lines.append("")
        return lines

    def _wrap(self, text: str, width: int = 100) -> list[str]:
        words = text.split()
        if not words:
            return [""]
        lines: list[str] = []
        current = words[0]
        for word in words[1:]:
            if len(current) + 1 + len(word) <= width:
                current = f"{current} {word}"
            else:
                lines.append(current)
                current = word
        lines.append(current)
        return lines

    def _file_record(self, target: str, path: Path) -> dict[str, Any]:
        return {
            "target": target,
            "path": str(path.resolve()),
            "exists": path.exists(),
            "size_bytes": path.stat().st_size if path.exists() else 0,
        }
