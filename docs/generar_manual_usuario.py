#!/usr/bin/env python3
"""Genera el Manual de Usuario ARIHAX Dx Pro — diseño y tono alineados al producto."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "MANUAL_USUARIO_ARHIAX_Dx_Pro_v1.1.docx"
APP_URL = "https://arhiax-dx-pro-187668243215.southamerica-east1.run.app"

# Paleta ARIHAX Dx Pro (misma que la UI)
C = {
    "cream": RGBColor(0xF4, 0xF1, 0xEA),
    "ink": RGBColor(0x17, 0x17, 0x17),
    "muted": RGBColor(0x70, 0x6F, 0x69),
    "green": RGBColor(0x56, 0x62, 0x4B),
    "green_light": RGBColor(0x78, 0x81, 0x5D),
    "bronze": RGBColor(0x9B, 0x6D, 0x4D),
    "navy": RGBColor(0x24, 0x3C, 0x4F),
    "error": RGBColor(0x8B, 0x3A, 0x3A),
    "rail": RGBColor(0x1A, 0x1B, 0x19),
    "white": RGBColor(0xFF, 0xFF, 0xFF),
}

FONT_TITLE = "Georgia"       # sustituto de Cormorant Garamond
FONT_BODY = "Calibri"        # sustituto de Manrope
FONT_MONO = "Consolas"       # sustituto de IBM Plex Mono


def _set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_font(run, name: str, size: Pt | None = None, color: RGBColor | None = None, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = C["ink"]
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    for level, size in [(1, 22), (2, 16), (3, 13)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = FONT_TITLE
        h.font.size = Pt(size)
        h.font.color.rgb = C["ink"]
        h.font.bold = False
        h._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_TITLE)


def add_mono_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    _set_run_font(run, FONT_MONO, Pt(9), C["bronze"], bold=True)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(12)


def add_body(doc: Document, text: str, space_after: int = 8) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    for run in p.runs:
        _set_run_font(run, FONT_BODY, Pt(11), C["ink"])


def add_callout(doc: Document, title: str, body: str, kind: str = "info") -> None:
    colors = {
        "info": ("E8EDE4", C["green"], C["green"]),
        "success": ("E8EDE4", C["green"], C["green"]),
        "warning": ("F5EDE6", C["bronze"], C["bronze"]),
        "error": ("F3E4E4", C["error"], C["error"]),
        "next": ("E8EDE4", C["green"], C["green"]),
    }
    bg, border, title_color = colors.get(kind, colors["info"])

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, bg)

    p0 = cell.paragraphs[0]
    r0 = p0.add_run(title)
    _set_run_font(r0, FONT_MONO, Pt(9), title_color, bold=True)
    p0.paragraph_format.space_after = Pt(4)

    p1 = cell.add_paragraph(body)
    p1.paragraph_format.line_spacing = 1.2
    for run in p1.runs:
        _set_run_font(run, FONT_BODY, Pt(10.5), C["ink"])

    # borde izquierdo simulado con párrafo vacío — Word no tiene border-left fácil en celdas sin XML extra
    doc.add_paragraph()


def add_image_placeholder(doc: Document, fig_id: str, caption: str, height_cm: float = 8.5) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "F4F1EA")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r1 = p.add_run("┌─────────────────────────────────────────┐\n")
    _set_run_font(r1, FONT_MONO, Pt(9), C["muted"])
    r2 = p.add_run(f"  {fig_id}\n  INSERTAR CAPTURA DE PANTALLA AQUÍ\n")
    _set_run_font(r2, FONT_MONO, Pt(10), C["muted"], italic=True)
    r3 = p.add_run("└─────────────────────────────────────────┘")
    _set_run_font(r3, FONT_MONO, Pt(9), C["muted"])

    cap = cell.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        _set_run_font(run, FONT_BODY, Pt(9.5), C["muted"], italic=True)

    cell.height = Cm(height_cm)
    doc.add_paragraph()


def add_steps(doc: Document, steps: list[str]) -> None:
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(step)
        _set_run_font(run, FONT_BODY, Pt(11), C["ink"])


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, FONT_BODY, Pt(11), C["ink"])


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_bg: str = "1A1B19") -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, FONT_MONO, Pt(9), C["white"], bold=True)
        _set_cell_shading(cell, header_bg)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            _set_run_font(run, FONT_BODY, Pt(10), C["ink"])
            if i % 2 == 0:
                _set_cell_shading(cell, "F4F1EA")
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    # banda superior oscura (rail)
    band = doc.add_table(rows=1, cols=1)
    cell = band.rows[0].cells[0]
    _set_cell_shading(cell, "1A1B19")
    cell.height = Cm(3.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("  Dx")
    _set_run_font(r, FONT_MONO, Pt(14), C["white"], bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Manual de operación")
    _set_run_font(r, FONT_MONO, Pt(11), C["bronze"], bold=True)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("ARHIAX Dx Pro")
    _set_run_font(r2, FONT_TITLE, Pt(36), C["ink"])

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("Diagnóstico organizacional gobernado · PMEL/ATK")
    _set_run_font(r3, FONT_BODY, Pt(12), C["muted"])

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(line.add_run("—"), FONT_BODY, Pt(14), C["muted"])

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    txt = (
        f"Versión 1.1 · {date.today().strftime('%d de %B de %Y')}\n"
        f"Entorno: {APP_URL}\n"
        "Audiencia: consultores y validadores HIL"
    )
    _set_run_font(meta.add_run(txt), FONT_MONO, Pt(10), C["muted"])

    doc.add_page_break()


def build() -> Path:
    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    # ── Índice ───────────────────────────────────────────────────────────────
    doc.add_heading("Índice", level=1)
    add_body(doc, "Actualice en Word: Referencias → Tabla de contenido → Actualizar tabla.")
    for item in [
        "1. Para quién es este manual y qué resuelve",
        "2. Acceso, credenciales y primer ingreso",
        "3. Mapa de pantallas",
        "4. Operación estándar — de cero a informe PDF",
        "5. Crear un caso (wizard paso a paso)",
        "6. Gestionar un caso en cada estado",
        "7. Encuesta pública — instrucciones para participantes",
        "8. Validación HIL y entregables",
        "9. Casos de error — qué significan y qué hacer",
        "10. Ejemplos reales del entorno (referencia operativa)",
        "11. Clientes, revisiones, evidencia y compliance",
        "12. Glosario",
        "13. Matriz de resolución de incidentes",
        "Anexo A — Checklist de capturas",
        "Anexo B — Rutas URL del sistema",
    ]:
        add_body(doc, item, space_after=2)
    doc.add_page_break()

    # ── 1 ────────────────────────────────────────────────────────────────────
    doc.add_heading("1. Para quién es este manual y qué resuelve", level=1)
    add_body(doc,
        "Este documento describe cómo opera ARIHAX Dx Pro en el día a día: "
        "crear diagnósticos, distribuir encuestas multi-rater, ejecutar el pipeline "
        "de IA bajo gobernanza PMEL/ATK, validar resultados (HIL) y entregar el informe. "
        "No es documentación técnica de desarrollo; es la guía del consultor que usa la plataforma."
    )
    add_callout(doc, "PRINCIPIO OPERATIVO",
        "El sistema no inventa datos si algo falla (fail-closed). "
        "Si un caso queda en Error, no hay PDF parcial ni atajos: se corrige la causa o se crea un caso nuevo.",
        "warning")
    add_body(doc, "Flujo resumido en cuatro movimientos:")
    add_table(doc, ["Fase", "Qué ocurre", "Quién actúa"], [
        ["1. Configuración", "Wizard → arquitectura G09 del instrumento", "Consultor"],
        ["2. Recolección", "Participantes responden por rol", "Cliente / informantes"],
        ["3. Fusión", "Pipeline IA procesa respuestas reales", "Sistema (automático)"],
        ["4. Cierre", "Revisión HIL → aprobación → PDF", "Consultor senior"],
    ])
    add_image_placeholder(doc, "Fig. 1.1", "Vista general del dashboard Pro con menú lateral y listado de casos")

    # ── 2 ────────────────────────────────────────────────────────────────────
    doc.add_heading("2. Acceso, credenciales y primer ingreso", level=1)
    add_mono_label(doc, "§ acceso")
    add_table(doc, ["Recurso", "Valor"], [
        ["URL login", f"{APP_URL}/login"],
        ["URL dashboard", f"{APP_URL}/dashboard-pro"],
        ["Usuario demo consultor", "admin@arhiax.com"],
        ["Contraseña demo", "arhiax-admin-2026"],
    ])
    add_steps(doc, [
        "Abra /login en el navegador (Chrome o Edge recomendados).",
        "Ingrese correo y contraseña. Pulse Iniciar sesión.",
        "Será redirigido a /dashboard-pro.",
        "Si la sesión expiró (mensaje en pantalla o error al guardar), cierre y vuelva a entrar — no pierde casos ya creados.",
    ])
    add_callout(doc, "SESIÓN EXPIRADA",
        "Síntoma: al crear un caso aparece error de conexión o redirección al login. "
        "Causa: token JWT vencido. Acción: reingresar y continuar desde el listado de casos.",
        "error")
    add_image_placeholder(doc, "Fig. 2.1", "Pantalla de login — campos, botón, logo ARHIAX")
    add_image_placeholder(doc, "Fig. 2.2", "Mensaje de sesión expirada (si aplica)")

    # ── 3 ────────────────────────────────────────────────────────────────────
    doc.add_heading("3. Mapa de pantallas", level=1)
    add_table(doc, ["Menú", "Ruta", "Uso"], [
        ["Casos", "/dashboard-pro", "Listado y filtros de todos los diagnósticos"],
        ["Nuevo caso", "/dashboard-pro/new", "Wizard de creación"],
        ["Clientes", "/dashboard-pro/clients", "Casos agrupados por organización"],
        ["Revisiones", "/dashboard-pro/reviews", "Cola HIL y decisiones recientes"],
        ["Evidencia", "/dashboard-pro/evidence", "Ledger global de eventos"],
        ["Compliance", "/dashboard-pro/compliance", "Postura PMEL/ATK del runtime"],
        ["Detalle caso", "/dashboard-pro/cases/{uuid}", "Operación central de un diagnóstico"],
        ["Encuesta pública", "/survey/pro/{token}", "Sin login — participantes"],
    ])
    add_image_placeholder(doc, "Fig. 3.1", "Menú lateral completo (fondo oscuro #1a1b19, ítems activos resaltados)")

    # ── 4 ────────────────────────────────────────────────────────────────────
    doc.add_heading("4. Operación estándar — de cero a informe PDF", level=1)
    add_body(doc, "Procedimiento que debe seguir todo consultor. Tiempo estimado: 2–5 días según respuesta de participantes.")
    add_steps(doc, [
        "Nuevo caso → completar Perfil (datos cliente + síntoma claro, máx. ~500 palabras).",
        "Arquitectura → seleccionar 3 roles (Estratégico, Operativo, Táctico), dimensiones e hipótesis con incidente real.",
        "Validación → revisar resumen y enviar. El caso pasa a Diseñando.",
        "Esperar arquitectura G09 (1–3 min). Estado cambia a Encuesta abierta.",
        "Copiar enlace del Hub de recolección. Enviar a un participante por rol.",
        "Verificar contador X/Y respuestas. Mínimo = cantidad de roles configurados.",
        "Lanzar síntesis de diagnóstico. Estado: Ejecutando (5–15 min).",
        "Al terminar: En revisión. Revisar paneles de resultados.",
        "Aprobar (HIL) → descargar PDF desde panel de entregables.",
    ])
    add_callout(doc, "PRÓXIMO PASO",
        "En el detalle de cada caso, el banner verde «Próximo paso» indica exactamente qué hacer ahora. "
        "Úselo como guía principal si no está seguro del estado.",
        "next")
    add_image_placeholder(doc, "Fig. 4.1", "Stepper del caso: Arquitectura → Recolección → Fusión IA → Validación HIL")

    # ── 5 Wizard ─────────────────────────────────────────────────────────────
    doc.add_heading("5. Crear un caso (wizard paso a paso)", level=1)
    add_mono_label(doc, "§ wizard · /dashboard-pro/new")

    doc.add_heading("5.1 Paso 1 — Perfil", level=2)
    add_body(doc, "Datos que alimentan el diseño del instrumento. El síntoma es crítico: debe ser concreto y medible.")
    add_callout(doc, "REGLA DEL SÍNTOMA",
        "Mínimo 20 caracteres. Recomendado: 2–4 oraciones (80–400 caracteres). "
        "Evite pegar informes completos: casos con +4.000 caracteres han fallado en arquitectura por límite de tokens.",
        "warning")
    add_bullets(doc, [
        "Identidad: nombre comercial, razón social, NIT, sector, ciudad, tamaño.",
        "Contacto: nombre, cargo, email, teléfono.",
        "Problema: dominio, síntoma, desde cuándo, resultado esperado.",
        "Alcance: engagement ID, confidencialidad.",
    ])
    add_image_placeholder(doc, "Fig. 5.1a", "Paso 1 — sección Identidad y Contacto", 7)
    add_image_placeholder(doc, "Fig. 5.1b", "Paso 1 — síntoma y dominio (campo El problema)", 7)

    doc.add_heading("5.2 Paso 2 — Arquitectura", level=2)
    add_body(doc, "Separe claramente roles (quién responde) de dimensiones (qué se evalúa).")
    add_table(doc, ["Concepto", "Opciones", "Notas"], [
        ["Roles (multi-rater)", "Estratégico · Operativo · Táctico", "3 perspectivas. Cada una responde una vez."],
        ["Dimensiones", "Estrategia, Procesos, Tecnología, Personas, Finanzas, Gobernanza", "Qué mide el instrumento, no quién responde."],
        ["Hipótesis", "Enunciado + incidente + informante", "Ancla las preguntas a hechos reales."],
    ])
    add_image_placeholder(doc, "Fig. 5.2a", "Selección de 3 roles con texto explicativo", 7)
    add_image_placeholder(doc, "Fig. 5.2b", "Dimensiones e hipótesis ancladas", 8)

    doc.add_heading("5.3 Paso 3 — Validación y envío", level=2)
    add_steps(doc, [
        "Revise tarjetas Empresa, El problema y Alcance.",
        "Marque consentimientos.",
        "Pulse Iniciar ciclo de diagnóstico.",
        "Será redirigido al detalle del caso en estado Diseñando.",
    ])
    add_image_placeholder(doc, "Fig. 5.3", "Resumen final y botón de envío")

    # ── 6 Estados ────────────────────────────────────────────────────────────
    doc.add_heading("6. Gestionar un caso en cada estado", level=1)
    add_mono_label(doc, "§ caso · /dashboard-pro/cases/{id}")

    add_table(doc, ["Estado en UI", "Significado", "Acción inmediata"], [
        ["Borrador", "Sin enviar", "Completar wizard"],
        ["Diseñando", "G09 generando preguntas", "Esperar — pantalla se actualiza sola"],
        ["Encuesta abierta", "Instrumento listo", "Compartir enlace; monitorear X/Y"],
        ["Ejecutando", "Pipeline IA", "No cerrar; esperar banner de fusión"],
        ["En revisión", "Listo para HIL", "Revisar y aprobar/rechazar"],
        ["Aprobado", "Sellado", "Descargar PDF"],
        ["Error", "Pipeline bloqueado", "Leer mensaje; ver capítulo 9"],
    ])

    doc.add_heading("6.1 Hub de recolección", level=2)
    add_steps(doc, [
        "Localice el bloque ENLACE SEGURO DE ACCESO.",
        "Pulse COPIAR y envíe por canal acordado con el cliente (email, Teams, etc.).",
        "Verifique PARTICIPANTES ESPERADOS — debe coincidir con los roles definidos.",
        "Cuando responses_count ≥ min_responses, habilite Lanzar síntesis de diagnóstico.",
    ])
    add_image_placeholder(doc, "Fig. 6.1", "Hub: enlace, contadores, chips de roles, botón lanzar")

    doc.add_heading("6.2 Fusión y revisión", level=2)
    add_image_placeholder(doc, "Fig. 6.2a", "Banner ciclo de fusión en ejecución + pipeline de etapas")
    add_image_placeholder(doc, "Fig. 6.2b", "Panel HIL lateral — Aprobar / Rechazar")

    doc.add_heading("6.3 Resultados y descarga", level=2)
    add_body(doc, "Tras aprobar: PDF, DOCX y MD según disponibilidad. El sello criptográfico queda en evidencia.")
    add_image_placeholder(doc, "Fig. 6.3", "Panel de entregables con botones de descarga")

    # ── 7 Encuesta ───────────────────────────────────────────────────────────
    doc.add_heading("7. Encuesta pública — instrucciones para participantes", level=1)
    add_body(doc, "Texto que puede copiar y enviar al cliente junto con el enlace:")
    add_callout(doc, "MENSAJE SUGERIDO AL CLIENTE",
        "«Estimado/a: le comparto el enlace del diagnóstico organizacional. "
        "Seleccione el rol que mejor describe su función (solo uno por persona). "
        "Tiempo estimado: 15–20 minutos. Sus respuestas son anónimas y tratadas bajo protocolo de gobernanza.»",
        "info")
    add_steps(doc, [
        "Abrir enlace /survey/pro/{token}.",
        "Leer introducción (organización, tiempo, privacidad).",
        "Elegir perspectiva (Estratégico / Operativo / Táctico).",
        "Responder ítems Likert (1–5) u abiertos.",
        "Finalizar — pantalla de confirmación.",
    ])
    add_image_placeholder(doc, "Fig. 7.1", "Selección de rol en encuesta pública")
    add_image_placeholder(doc, "Fig. 7.2", "Ítem Likert con escala 1–5")
    add_image_placeholder(doc, "Fig. 7.3", "Pantalla Diagnóstico Finalizado")

    # ── 8 HIL ─────────────────────────────────────────────────────────────────
    doc.add_heading("8. Validación HIL y entregables", level=1)
    add_body(doc, "La aprobación humana es obligatoria. Sin ella no hay descarga de informe.")
    add_steps(doc, [
        "Abra el caso en estado En revisión.",
        "Revise ProResultsPanel: madurez, dimensiones, riesgos, hipótesis.",
        "Opcional: SurveyAuditPanel para auditar respuestas por rol.",
        "Escriba observaciones si aplica.",
        "Pulse APROBAR — se genera sello y entregables.",
        "Descargue PDF y comparta con el cliente.",
    ])
    add_callout(doc, "RECHAZO",
        "Si rechaza, el caso queda en Rechazado. No hay PDF. Documente el motivo en observaciones para trazabilidad.",
        "warning")

    # ── 9 ERRORES ─────────────────────────────────────────────────────────────
    doc.add_heading("9. Casos de error — qué significan y qué hacer", level=1)
    add_body(doc,
        "Capítulo operativo crítico. Un caso en Error no se «arregla» con reintentos mágicos: "
        "hay que entender la causa y actuar según la tabla."
    )

    doc.add_heading("9.1 Error en arquitectura (G09)", level=2)
    add_table(doc, ["Síntoma en pantalla", "Causa probable", "Acción"], [
        ["Estado Error + encuesta no disponible", "Síntoma excesivamente largo (>~2.000 caracteres)", "Crear caso nuevo con síntoma resumido"],
        ["Encuesta: «no se generó correctamente»", "Fallo en generación de preguntas", "Revisar evidencia pipeline_failed; nuevo caso"],
        ["Survey status = error", "G09 no completó", "No compartir enlace; caso nuevo"],
    ])
    add_callout(doc, "EJEMPLO REAL — case-3046bf33de (CD Global)",
        "Cliente: CD Global. Síntoma: 4.577 caracteres (texto pegado completo). "
        "Resultado: pipeline_failed en agente G09, 0 preguntas, encuesta en error. "
        "Solución aplicada: no reutilizar ese caso; crear uno nuevo con síntoma de 2–4 párrafos.",
        "error")
    add_image_placeholder(doc, "Fig. 9.1", "Caso en error: banner rojo + mensaje pipeline + stepper con X")

    doc.add_heading("9.2 Error en fusión (post-recolección)", level=2)
    add_table(doc, ["Síntoma", "Causa probable", "Acción"], [
        ["Error tras Lanzar síntesis", "Etapa del pipeline sin salida LLM válida", "Leer pipeline_error; revisar evidencia"],
        ["No se puede aprobar (409)", "Etapas comprometidas", "No aprobar; caso queda bloqueado"],
        ["PDF no disponible", "Caso no aprobado o en error", "Completar HIL o corregir error primero"],
    ])
    add_callout(doc, "POLÍTICA FAIL-CLOSED",
        "El sistema no entrega informes con datos simulados ni omite etapas fallidas. "
        "Esto es intencional para auditoría y confianza del cliente.",
        "warning")

    doc.add_heading("9.3 Errores de operación del consultor", level=2)
    add_table(doc, ["Situación", "Qué ve", "Qué hacer"], [
        ["Sesión expirada", "Redirect a login o error al crear", "Reingresar; caso no se pierde"],
        ["Lanzar síntesis deshabilitado", "0 respuestas", "Esperar al menos 1 respuesta"],
        ["Faltan respuestas", "X/Y incompleto", "Recordar a participantes por rol faltante"],
        ["Rol duplicado en encuesta legacy", "4 roles incl. Planeación", "Cada rol distinto responde una vez"],
        ["Encuesta cerrada (410)", "Participante no puede entrar", "Caso ya pasó de fase; no reabrir"],
        ["PDF 500 en caso viejo", "Pre-pipeline sin validación", "Solo casos aprobados post-fix tienen PDF válido"],
    ])

    doc.add_heading("9.4 Qué NO hacer", level=2)
    add_bullets(doc, [
        "No pegar informes enteros en el campo síntoma.",
        "No mezclar dimensiones como si fueran roles (el wizard actual ya lo previene).",
        "No aprobar un caso sin revisar resultados — el sello HIL es su responsabilidad.",
        "No eliminar casos del entorno de producción sin acuerdo del equipo.",
        "No compartir el enlace de encuesta antes de que el estado sea Encuesta abierta.",
    ])

    # ── 10 Ejemplos reales ───────────────────────────────────────────────────
    doc.add_heading("10. Ejemplos reales del entorno (referencia operativa)", level=1)
    add_body(doc, "Casos existentes en producción al momento de redactar este manual. Sirven como referencia para el consultor senior.")

    add_table(doc, ["Referencia", "Cliente", "Estado", "Lección operativa"], [
        ["case-ffbf1ae77b", "Agencia Innovación São Caetano", "Aprobado", "Flujo completo OK: 3 respuestas, 21/21 etapas, PDF disponible"],
        ["case-b957403514", "CD GLOBAL (Ivania Rua)", "Encuesta abierta", "19 preguntas, 0/4 respuestas — pendiente distribuir enlace"],
        ["case-3046bf33de", "CD Global", "Error", "Síntoma 4.577 chars → fallo G09 — crear caso nuevo"],
        ["case-1b8c4eb343", "Smoke Test", "Encuesta abierta", "Prueba interna, 0/3 respuestas"],
    ])

    doc.add_heading("10.1 Recorrido exitoso — São Caetano", level=2)
    add_steps(doc, [
        "Caso creado con síntoma conciso (118 caracteres) y 3 hipótesis.",
        "Encuesta: 19 preguntas, 3 respuestas (una por rol).",
        "Síntesis lanzada → 21 etapas pipeline completadas.",
        "Aprobación HIL → 3 entregables generados.",
        "Índice de madurez: 50.0 — PDF descargable.",
    ])
    add_image_placeholder(doc, "Fig. 10.1", "Caso São Caetano aprobado — panel resultados y descarga PDF")

    doc.add_heading("10.2 Caso en curso — Ivania (CD Global)", level=2)
    add_body(doc, f"Enlace encuesta: {APP_URL}/survey/pro/01076dc4b8d24b129ec3f219f1185b26")
    add_steps(doc, [
        "Distribuir enlace a 4 participantes (roles: Estratégico, Operativo, Táctico, Planeación en casos legacy).",
        "Confirmar 4/4 respuestas en el hub.",
        "Lanzar síntesis.",
        "Revisar y aprobar.",
        "Entregar PDF al cliente.",
    ])
    add_image_placeholder(doc, "Fig. 10.2", "Caso Ivania — hub de recolección con 0/4 y roles esperados")

    doc.add_heading("10.3 Caso fallido — CD Global (síntoma largo)", level=2)
    add_body(doc, "Referencia de lo que NO debe repetirse. El consultor debe reconocer este patrón en el banner de error.")
    add_image_placeholder(doc, "Fig. 10.3", "Caso CD Global error — pipeline_failed visible en evidencia")

    # ── 11 Otras pantallas ───────────────────────────────────────────────────
    doc.add_heading("11. Clientes, revisiones, evidencia y compliance", level=1)

    doc.add_heading("11.1 Clientes", level=2)
    add_body(doc, "Agrupa diagnósticos por organización. Útil cuando un mismo cliente tiene varios engagements.")
    add_image_placeholder(doc, "Fig. 11.1", "Vista /dashboard-pro/clients")

    doc.add_heading("11.2 Revisiones", level=2)
    add_body(doc, "Cola de casos En revisión y historial de aprobados/rechazados.")
    add_image_placeholder(doc, "Fig. 11.2", "Vista /dashboard-pro/reviews")

    doc.add_heading("11.3 Evidencia", level=2)
    add_body(doc, "Ledger append-only. Filtro por trace_id. Botón Verificar HMAC valida integridad de la cadena.")
    add_image_placeholder(doc, "Fig. 11.3", "Vista /dashboard-pro/evidence + resultado HMAC")

    doc.add_heading("11.4 Compliance", level=2)
    add_body(doc, "Postura del runtime: policy engine, ledger HEAD, cobertura PMEL.")
    add_image_placeholder(doc, "Fig. 11.4", "Vista /dashboard-pro/compliance")

    # ── 12 Glosario ──────────────────────────────────────────────────────────
    doc.add_heading("12. Glosario", level=1)
    add_table(doc, ["Término", "Definición"], [
        ["Multi-rater", "Encuesta con varios roles; cada perspectiva responde por separado"],
        ["HIL", "Human-in-the-Loop — validación humana obligatoria antes de entregar"],
        ["PMEL/ATK", "Capa de políticas y gobernanza del pipeline"],
        ["G09", "Agentes que diseñan el instrumento de encuesta"],
        ["Fail-closed", "Ante fallo, el sistema se detiene sin datos ficticios"],
        ["Trace ID", "Identificador de trazabilidad en evidencia"],
        ["Hub de recolección", "Pantalla con enlace y contadores de respuestas"],
        ["Síntesis", "Ejecución del pipeline IA tras recolección"],
    ])

    # ── 13 Matriz incidentes ─────────────────────────────────────────────────
    doc.add_heading("13. Matriz de resolución de incidentes", level=1)
    add_table(doc, ["#", "Incidente", "Severidad", "Resolución", "Tiempo est."], [
        ["1", "Sesión expirada", "Baja", "Re-login", "< 1 min"],
        ["2", "Encuesta no disponible", "Media", "Esperar G09 o nuevo caso", "1–5 min"],
        ["3", "Caso Error G09", "Alta", "Nuevo caso, síntoma corto", "15 min"],
        ["4", "Fusión fallida", "Alta", "Revisar error; soporte si persiste", "Variable"],
        ["5", "PDF no descarga", "Media", "Verificar aprobación HIL", "Inmediato"],
        ["6", "Respuestas incompletas", "Baja", "Reenviar enlace a rol faltante", "Días"],
        ["7", "HMAC comprometido", "Crítica", "Escalar a administrador", "Urgente"],
    ])

    doc.add_page_break()

    # ── Anexo A ──────────────────────────────────────────────────────────────
    doc.add_heading("Anexo A — Checklist de capturas", level=1)
    add_body(doc, "Marque cada figura al insertarla. Total: 28 espacios.")
    figures = [
        "Fig. 1.1 Dashboard general", "Fig. 2.1 Login", "Fig. 2.2 Sesión expirada",
        "Fig. 3.1 Menú lateral", "Fig. 4.1 Stepper ciclo de vida",
        "Fig. 5.1a-b Wizard perfil", "Fig. 5.2a-b Wizard arquitectura", "Fig. 5.3 Validación",
        "Fig. 6.1 Hub recolección", "Fig. 6.2a Fusión", "Fig. 6.2b HIL", "Fig. 6.3 Entregables",
        "Fig. 7.1-7.3 Encuesta pública (3)", "Fig. 9.1 Caso error",
        "Fig. 10.1 São Caetano", "Fig. 10.2 Ivania", "Fig. 10.3 CD Global error",
        "Fig. 11.1-11.4 Clientes, revisiones, evidencia, compliance",
    ]
    for fig in figures:
        p = doc.add_paragraph(style="List Bullet")
        _set_run_font(p.add_run(f"{fig}   ☐"), FONT_MONO, Pt(10), C["ink"])

    doc.add_heading("Anexo B — Rutas URL", level=1)
    add_table(doc, ["Pantalla", "URL"], [
        ["Login", f"{APP_URL}/login"],
        ["Casos", f"{APP_URL}/dashboard-pro"],
        ["Nuevo caso", f"{APP_URL}/dashboard-pro/new"],
        ["Encuesta Ivania (ejemplo)", f"{APP_URL}/survey/pro/01076dc4b8d24b129ec3f219f1185b26"],
    ])

    # pie de documento
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run("— ARHIAX Dx Pro · Documento interno de operación —"),
                  FONT_MONO, Pt(9), C["muted"], italic=True)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Manual generado: {path}")
