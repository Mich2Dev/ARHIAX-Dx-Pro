"""
Document upload endpoints — optional context files for a diagnostic.
Supports: PDF, DOCX, TXT, XLSX, PNG, JPG (max 10 MB each, max 5 per diagnostic).
"""

from __future__ import annotations

import hashlib
import io
import logging
import uuid
from datetime import datetime, timezone
from pathlib import Path

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from api.auth import get_current_user
from api.db import get_db
from api.models import Diagnostic, DiagnosticDocument, User

log = logging.getLogger("arhiax.documents")

router = APIRouter(prefix="/v2/diagnostics", tags=["documents"])

# ── Config ────────────────────────────────────────────────────────────────────
MAX_FILE_SIZE   = 10 * 1024 * 1024   # 10 MB
MAX_FILES       = 5
UPLOAD_DIR      = Path("uploads/documents")
ALLOWED_TYPES   = {
    "application/pdf":                                                    "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword":                                                 "doc",
    "text/plain":                                                         "txt",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel":                                           "xls",
    "image/png":                                                          "png",
    "image/jpeg":                                                         "jpg",
}

DOC_TYPE_LABELS = {
    "context":      "Descripción del problema",
    "organigrama":  "Organigrama",
    "financiero":   "Datos financieros",
    "proceso":      "Manual de proceso",
    "otro":         "Otro",
}


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post("/{diagnostic_id}/documents")
async def upload_document(
    diagnostic_id: str,
    file: UploadFile = File(...),
    doc_type: str = Form(default="context"),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
) -> dict:
    """Upload an optional context document for a diagnostic."""

    # Validate diagnostic exists
    diagnostic = await db.get(Diagnostic, diagnostic_id)
    if not diagnostic:
        raise HTTPException(status_code=404, detail="Diagnostic not found")

    # Check file count limit
    existing = await db.execute(
        select(DiagnosticDocument)
        .where(DiagnosticDocument.diagnostic_id == diagnostic_id)
    )
    count = len(existing.scalars().all())
    if count >= MAX_FILES:
        raise HTTPException(
            status_code=400,
            detail=f"Maximum {MAX_FILES} documents per diagnostic"
        )

    # Validate mime type
    mime = file.content_type or ""
    if mime not in ALLOWED_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed. Supported: PDF, DOCX, TXT, XLSX, PNG, JPG"
        )

    # Read and validate size
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size: 10 MB"
        )

    # Save file to disk
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    ext = ALLOWED_TYPES[mime]
    stored_name = f"{diagnostic_id}_{uuid.uuid4().hex[:8]}.{ext}"
    file_path = UPLOAD_DIR / stored_name
    file_path.write_bytes(content)

    # Extract text for LLM context
    extracted_text = _extract_text(content, mime, file.filename or "")

    # Create DB record
    doc = DiagnosticDocument(
        diagnostic_id=diagnostic_id,
        filename=stored_name,
        original_name=file.filename or stored_name,
        mime_type=mime,
        size_bytes=len(content),
        doc_type=doc_type if doc_type in DOC_TYPE_LABELS else "otro",
        extracted_text=extracted_text,
        uploaded_by=user.id,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    log.info("Document uploaded: %s (%d bytes) for diagnostic %s", stored_name, len(content), diagnostic_id)

    return {
        "id": doc.id,
        "original_name": doc.original_name,
        "doc_type": doc.doc_type,
        "doc_type_label": DOC_TYPE_LABELS.get(doc.doc_type, "Otro"),
        "size_bytes": doc.size_bytes,
        "size_human": _human_size(doc.size_bytes),
        "has_text": bool(extracted_text),
        "created_at": doc.created_at.isoformat(),
    }


@router.get("/{diagnostic_id}/documents")
async def list_documents(
    diagnostic_id: str,
    db: AsyncSession = Depends(get_db),
    _user: User = Depends(get_current_user),
) -> dict:
    """List all documents for a diagnostic."""
    result = await db.execute(
        select(DiagnosticDocument)
        .where(DiagnosticDocument.diagnostic_id == diagnostic_id)
        .order_by(DiagnosticDocument.created_at.asc())
    )
    docs = result.scalars().all()
    return {
        "items": [_doc_out(d) for d in docs],
        "count": len(docs),
        "max": MAX_FILES,
    }


@router.delete("/{diagnostic_id}/documents/{document_id}")
async def delete_document(
    diagnostic_id: str,
    document_id: str,
    db: AsyncSession = Depends(get_db),
    _user: User = Depends(get_current_user),
) -> dict:
    """Delete a document."""
    result = await db.execute(
        select(DiagnosticDocument)
        .where(DiagnosticDocument.id == document_id)
        .where(DiagnosticDocument.diagnostic_id == diagnostic_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Delete file from disk
    file_path = UPLOAD_DIR / doc.filename
    if file_path.exists():
        file_path.unlink()

    await db.delete(doc)
    await db.commit()
    return {"success": True}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _extract_text(content: bytes, mime: str, filename: str) -> str | None:
    """Extract plain text from uploaded file for LLM context."""
    try:
        if mime == "text/plain":
            return content.decode("utf-8", errors="ignore")[:8000]

        if mime == "application/pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(content))
                text = "\n".join(
                    page.extract_text() or "" for page in reader.pages
                )
                return text[:8000] if text.strip() else None
            except Exception:
                return None

        if mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            try:
                from docx import Document
                doc = Document(io.BytesIO(content))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                return text[:8000] if text.strip() else None
            except Exception:
                return None

        if mime in (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
        ):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
                lines = []
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        line = " | ".join(str(c) for c in row if c is not None)
                        if line.strip():
                            lines.append(line)
                text = "\n".join(lines)
                return text[:8000] if text.strip() else None
            except Exception:
                return None

    except Exception as e:
        log.warning("Text extraction failed for %s: %s", filename, e)

    return None


def _doc_out(d: DiagnosticDocument) -> dict:
    return {
        "id": d.id,
        "original_name": d.original_name,
        "doc_type": d.doc_type,
        "doc_type_label": DOC_TYPE_LABELS.get(d.doc_type, "Otro"),
        "mime_type": d.mime_type,
        "size_bytes": d.size_bytes,
        "size_human": _human_size(d.size_bytes),
        "has_text": bool(d.extracted_text),
        "created_at": d.created_at.isoformat() if d.created_at else None,
    }


def _human_size(size: int) -> str:
    if size < 1024:
        return f"{size} B"
    if size < 1024 * 1024:
        return f"{size / 1024:.1f} KB"
    return f"{size / (1024 * 1024):.1f} MB"
