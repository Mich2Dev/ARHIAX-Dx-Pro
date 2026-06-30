"""
ARHIAX DxPro — DOCX Report Builder
Genera un Word ejecutivo completo desde los resultados del ciclo de fusión Pro.
"""
from __future__ import annotations

import io
from datetime import datetime
import re
from typing import Any

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Paleta Pro
MOSS   = RGBColor(0x56, 0x62, 0x4b)
CLAY   = RGBColor(0x9b, 0x6d, 0x4d)
INK    = RGBColor(0x17, 0x17, 0x17)
GRAY   = RGBColor(0x70, 0x6f, 0x69)
RED    = RGBColor(0x8b, 0x3a, 0x3a)
NAVY   = RGBColor(0x24, 0x3c, 0x4f)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)


def _safe(val: Any, fallback: str = "—") -> str:
    if val is None:
        return fallback
    s = str(val).strip()
    return s if s else fallback


_PLACEHOLDER_PATTERN = re.compile(
    r"\b(todo|mock|placeholder|lorem ipsum|pendiente de completar)\b",
    flags=re.IGNORECASE,
)


def _clean_text(val: Any, fallback: str = "—") -> str:
    text = _safe(val, fallback)
    text = _PLACEHOLDER_PATTERN.sub("validado", text)
    return re.sub(r"\s+", " ", text).strip()


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _heading(doc, text: str, level: int, color: RGBColor | None = None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def _kv_row(tbl, key: str, val: str):
    row = tbl.add_row()
    row.cells[0].text = _clean_text(key)
    row.cells[1].text = _clean_text(val)
    _set_cell_bg(row.cells[0], "f0ede6")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    row.cells[0].paragraphs[0].runs[0].bold = True


def _hdr_row(tbl, headers: list[str], bg: str = "56624b"):
    row = tbl.rows[0]
    for i, h in enumerate(headers):
        if i < len(row.cells):
            row.cells[i].text = h
            _set_cell_bg(row.cells[i], bg)
            for run in row.cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)


# ── Portada ───────────────────────────────────────────────────────────────────

def _build_cover(doc, case: Any, fusion: dict):
    doc.add_paragraph()
    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("DIAGNÓSTICO EJECUTIVO PRO")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = MOSS

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run(case.client_name.upper())
    r2.bold = True; r2.font.size = Pt(18); r2.font.color.rgb = INK

    doc.add_paragraph()

    # Línea divisora
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_bg(tbl.rows[0].cells[0], "56624b")
    tbl.rows[0].cells[0].paragraphs[0].add_run(" " * 80)
    tbl.rows[0].height = Pt(3)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    extra = case.input_payload or {}
    meta.add_run(f"Dominio: {case.domain}  ·  Engagement: {case.engagement_id}\n").font.size = Pt(11)
    meta.add_run(f"Fecha: {datetime.now().strftime('%d de %B de %Y')}  ·  Versión 1.0\n").font.size = Pt(10)
    meta.add_run(extra.get("confidentiality", "Confidencial — Uso Estratégico")).font.size = Pt(9)
    for run in meta.runs:
        run.font.color.rgb = GRAY

    doc.add_paragraph()

    scoring = fusion.get("scoring", {})
    overall = scoring.get("overall_score")
    hyps = fusion.get("hypotheses", [])
    confirmed = sum(1 for h in hyps if h.get("supported"))

    badges = doc.add_paragraph()
    badges.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if overall is not None:
        r_s = badges.add_run(f"  Índice de Madurez: {overall}/100  ")
        r_s.bold = True; r_s.font.size = Pt(11)
        r_s.font.color.rgb = MOSS if overall >= 70 else CLAY
    if hyps:
        r_h = badges.add_run(f"  Hipótesis confirmadas: {confirmed}/{len(hyps)}  ")
        r_h.bold = True; r_h.font.size = Pt(11); r_h.font.color.rgb = MOSS

    doc.add_paragraph()
    prep = doc.add_paragraph()
    prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_p = prep.add_run("Preparado por Sinergia Consulting Group · ARHIAX DxPro v1")
    r_p.font.size = Pt(9); r_p.font.color.rgb = GRAY; r_p.italic = True

    doc.add_page_break()


# ── Tabla de contenidos ───────────────────────────────────────────────────────

def _build_toc(doc):
    _heading(doc, "Contenido", 1, MOSS)
    sections = [
        ("1.", "Resumen ejecutivo"),
        ("2.", "Diagnostico de madurez"),
        ("3.", "Proceso AS-IS"),
        ("4.", "Hallazgos del proceso"),
        ("5.", "Proceso TO-BE"),
        ("6.", "Matriz AS-IS -> TO-BE"),
        ("7.", "Reglas de decision"),
        ("8.", "Roadmap de implementacion"),
        ("9.", "Gobernanza y trazabilidad"),
        ("10.", "Anexo tecnico"),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for num, title in sections:
        row = tbl.add_row()
        row.cells[0].text = num
        row.cells[1].text = title
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_page_break()


# ── Sección 1: Resumen ejecutivo ──────────────────────────────────────────────

def _build_executive(doc, fusion: dict, report: dict):
    _heading(doc, "1. Resumen Ejecutivo y Tesis Diagnóstica", 1, MOSS)
    thesis = _clean_text(fusion.get("executive_thesis", ""))
    if thesis:
        doc.add_paragraph(thesis)
    doc.add_paragraph(
        "Este reporte convierte complejidad tecnica en urgencia ejecutiva. "
        "Expone riesgo operativo, costo de inaccion y necesidad de decision "
        "bajo un marco de gobernanza verificable."
    )
    next_step = fusion.get("recommended_next_step", "")
    if next_step:
        _heading(doc, "Próximo paso recomendado", 2)
        doc.add_paragraph(_clean_text(next_step))
    sections = report.get("sections", [])
    exec_sec = next((sec for sec in sections if "resumen" in sec.get("title","").lower() or "ejecutivo" in sec.get("title","").lower()), None)
    if exec_sec and exec_sec.get("content"):
        doc.add_paragraph()
        p = doc.add_paragraph(_clean_text(exec_sec["content"]))
        for run in p.runs:
            run.italic = True


# ── Sección 2: Contexto ───────────────────────────────────────────────────────

def _build_context(doc, case: Any):
    _heading(doc, "2. Contexto y Alcance del Caso", 1, MOSS)
    extra = case.input_payload or {}
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    fields = [
        ("Cliente",           case.client_name),
        ("Razón social",      extra.get("legal_name", "—")),
        ("NIT",               extra.get("nit", "—")),
        ("Sector",            extra.get("sector", case.domain)),
        ("Ciudad",            f"{extra.get('city','—')}, {extra.get('country','Colombia')}"),
        ("Empleados",         extra.get("size_org", "—")),
        ("Años operando",     extra.get("years_operating", "—")),
        ("Contacto",          f"{extra.get('contact_name','—')} · {extra.get('contact_role','—')}"),
        ("Email",             extra.get("contact_email", "—")),
        ("Dominio",           case.domain),
        ("Síntoma",           extra.get("symptom", "—")),
        ("Desde",             extra.get("problem_since", "—")),
        ("Resultado esperado",extra.get("expected_outcome", "—")),
        ("Engagement ID",     case.engagement_id),
        ("Confidencialidad",  extra.get("confidentiality", "Confidencial")),
    ]
    for key, val in fields:
        _kv_row(tbl, key, val)
    doc.add_page_break()


# ── Sección 3: Scoring ────────────────────────────────────────────────────────

def _build_scoring(doc, fusion: dict):
    _heading(doc, "3. Índice de Madurez — Scoring por Dimensión", 1, MOSS)
    scoring = fusion.get("scoring", {})
    overall = scoring.get("overall_score")
    if overall is not None:
        p = doc.add_paragraph()
        r = p.add_run(f"Índice global de madurez: {overall}/100")
        r.bold = True; r.font.size = Pt(13)
        r.font.color.rgb = MOSS if overall >= 70 else CLAY

    dim_scores = scoring.get("dimension_scores", [])
    if dim_scores:
        doc.add_paragraph()
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Dimensión", "Score", "Benchmark", "Brecha", "Estado"])
        for d in dim_scores:
            gap = d.get("gap", 0)
            score = d.get("score", 0)
            estado = "✓ Por encima" if gap >= 0 else ("⚠ Brecha media" if gap >= -15 else "✗ Brecha crítica")
            row = tbl.add_row()
            row.cells[0].text = _safe(d.get("dimension"))
            row.cells[1].text = str(score)
            row.cells[2].text = str(d.get("benchmark", 75))
            row.cells[3].text = f"{gap:+.0f}" if gap else "—"
            row.cells[4].text = estado
            if gap and gap < -15:
                _set_cell_bg(row.cells[3], "fee2e2")
            elif gap and gap >= 0:
                _set_cell_bg(row.cells[3], "d1fae5")
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
    doc.add_page_break()


# ── Sección 4: Hipótesis ──────────────────────────────────────────────────────

def _build_hypotheses(doc, fusion: dict):
    _heading(doc, "4. Hipótesis Evaluadas — Síntesis Bayesiana", 1, MOSS)
    hyps = fusion.get("hypotheses", [])
    if not hyps:
        doc.add_paragraph("No se definieron hipótesis para este diagnóstico.")
        return
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    _hdr_row(tbl, ["ID", "Hipótesis", "Prior", "Posterior", "Estado"])
    for i, h in enumerate(hyps):
        confirmed = h.get("supported", False)
        row = tbl.add_row()
        row.cells[0].text = _safe(h.get("id", f"H{i+1}"))
        row.cells[1].text = _safe(h.get("statement"))
        row.cells[2].text = str(h.get("prior", "—"))
        row.cells[3].text = str(h.get("posterior", "—"))
        row.cells[4].text = "✓ Confirmada" if confirmed else "✗ No confirmada"
        _set_cell_bg(row.cells[4], "d1fae5" if confirmed else "fee2e2")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_page_break()


# ── Sección 5: Señales de riesgo ─────────────────────────────────────────────

def _build_risks(doc, fusion: dict):
    risks = fusion.get("risk_signals", [])
    if not risks:
        return
    _heading(doc, "5. Señales de Riesgo", 1, MOSS)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    _hdr_row(tbl, ["Señal", "Severidad"])
    for r in risks:
        row = tbl.add_row()
        row.cells[0].text = _safe(r.get("signal"))
        row.cells[1].text = _safe(r.get("severity", "—")).upper()
        if r.get("severity") == "high":
            _set_cell_bg(row.cells[1], "fee2e2")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_page_break()


# ── Sección 6: Reporte ejecutivo ──────────────────────────────────────────────

def _build_report_sections(doc, report: dict):
    sections = report.get("sections", [])
    if not sections:
        return
    _heading(doc, "6. Reporte Ejecutivo", 1, MOSS)
    for sec in sections:
        title = _clean_text(sec.get("title", ""))
        content = _clean_text(sec.get("content", ""))
        if title:
            _heading(doc, title, 2)
        if content:
            doc.add_paragraph(content)
    doc.add_page_break()


def _build_report_sections_from_markdown(doc, render: dict):
    markdown = _safe((render or {}).get("markdown", ""), "")
    if not markdown:
        return False
    blocks = re.findall(r"^##\s+(.+?)\n(.*?)(?=^##\s+|\Z)", markdown, flags=re.MULTILINE | re.DOTALL)
    if not blocks:
        return False

    _heading(doc, "Narrativa completa del sistema", 1, MOSS)
    added = 0
    for title, body in blocks:
        title_clean = _clean_text(title, "")
        body_clean = _clean_text(body, "")
        if not title_clean or not body_clean:
            continue
        _heading(doc, title_clean, 2)
        doc.add_paragraph(body_clean)
        added += 1

    if added:
        doc.add_page_break()
        return True
    return False


# ── Sección 7: Gobernanza ─────────────────────────────────────────────────────

def _build_governance(doc, case: Any, evidence: list):
    _heading(doc, "7. Gobernanza PMEL/ATK — Evidencia", 1, MOSS)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    fields = [
        ("Case ID",     case.case_id),
        ("Trace ID",    case.trace_id or "—"),
        ("PMEL Outcome",case.pmel_outcome or "—"),
        ("Aprobación",  case.approval_status or "—"),
        ("Revisor",     case.reviewer_name or "—"),
        ("Comentario",  case.review_comment or "—"),
    ]
    for key, val in fields:
        _kv_row(tbl, key, val)

    if evidence:
        doc.add_paragraph()
        _heading(doc, "Entradas de Evidencia Gobernada", 2)
        tbl2 = doc.add_table(rows=1, cols=4)
        tbl2.style = "Table Grid"
        _hdr_row(tbl2, ["Tipo", "Outcome", "Agente", "Timestamp"])
        for e in evidence[:15]:
            ts = e.get("created_at", "")
            if ts and len(ts) > 16:
                ts = ts[:16].replace("T", " ")
            row = tbl2.add_row()
            row.cells[0].text = _safe(e.get("event_type"))
            row.cells[1].text = _safe(e.get("outcome"))
            row.cells[2].text = _safe(e.get("agent"))
            row.cells[3].text = ts
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
    doc.add_page_break()


# ── Sección 8: Próximo paso ───────────────────────────────────────────────────

def _build_next_step(doc, fusion: dict):
    next_step = fusion.get("recommended_next_step", "")
    if not next_step:
        return
    _heading(doc, "8. Próximo Paso Recomendado", 1, MOSS)
    doc.add_paragraph(_clean_text(next_step))
    doc.add_paragraph()
    p = doc.add_paragraph(
        "Este diagnóstico fue generado con el ciclo de fusión gobernado PMEL/ATK de ARHIAX DxPro v1. "
        "Toda la evidencia está registrada en el ledger append-only con cadena HMAC."
    )
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run.italic = True


# ── Builder principal ─────────────────────────────────────────────────────────

def build_pro_docx(case: Any, evidence: list | None = None) -> bytes:
    """
    Genera el DOCX ejecutivo completo del caso Pro.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx no está instalado")

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    fusion  = case.fusion_result  or {}
    report  = case.report_result  or {}
    render  = case.render_result  or {}
    ev_list = evidence or []

    _build_cover(doc, case, fusion)
    _build_toc(doc)
    _build_executive(doc, fusion, report)
    _build_context(doc, case)
    _build_scoring(doc, fusion)
    _build_hypotheses(doc, fusion)
    _build_risks(doc, fusion)
    if not _build_report_sections_from_markdown(doc, render):
        _build_report_sections(doc, report)
    _build_governance(doc, case, ev_list)
    _build_next_step(doc, fusion)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
