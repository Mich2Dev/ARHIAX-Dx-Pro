"""
ARHIAX Dx — Executive Report Builder (DOCX)
Generates a professional Word document from pipeline outputs.
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Brand colors
GREEN  = RGBColor(0x0A, 0x7F, 0x5A)   # Sinergia brand
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
RED    = RGBColor(0xDC, 0x26, 0x26)
ORANGE = RGBColor(0xEA, 0x58, 0x0C)
BLUE   = RGBColor(0x25, 0x63, 0xEB)
PURPLE = RGBColor(0x7C, 0x3A, 0xED)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _heading(doc: Any, text: str, level: int, color: RGBColor | None = None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def _add_kv_row(table, key: str, value: str, shade_key: bool = True):
    row = table.add_row()
    row.cells[0].text = key
    row.cells[1].text = str(value)
    if shade_key:
        _set_cell_bg(row.cells[0], "F0FDF4")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].bold = True


def _score_bar_text(score: int, max_score: int = 100) -> str:
    """ASCII progress bar for score."""
    filled = int((score / max_score) * 20)
    bar = "█" * filled + "░" * (20 - filled)
    return f"{bar}  {score}/{max_score}"


def _safe_str(val: Any, fallback: str = "—") -> str:
    if val is None:
        return fallback
    s = str(val).strip()
    return s if s else fallback


def _extract(outputs: dict, *keys: str) -> dict:
    """Try multiple key names, return first non-empty dict."""
    for k in keys:
        v = outputs.get(k, {})
        if v:
            return v if isinstance(v, dict) else {}
    return {}


# ── Cover Page ────────────────────────────────────────────────────────────────

def _build_cover(doc: Any, diagnostic: Any, qa: dict, irr: dict):
    doc.add_paragraph()
    doc.add_paragraph()

    # Main title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("DIAGNÓSTICO ORGANIZACIONAL")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = GREEN

    # Org name
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run(diagnostic.organization_name.upper())
    r2.bold = True
    r2.font.size = Pt(20)
    r2.font.color.rgb = DARK

    doc.add_paragraph()

    # Divider line via table
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_bg(tbl.rows[0].cells[0], "0A7F5A")
    tbl.rows[0].cells[0].paragraphs[0].add_run(" " * 80)
    tbl.rows[0].height = Pt(4)

    doc.add_paragraph()

    # Metadata block
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Sector: {diagnostic.domain}  ·  Área: {diagnostic.subprocess}\n").font.size = Pt(12)
    meta.add_run(f"Fecha: {datetime.now().strftime('%d de %B de %Y')}  ·  Versión 1.0\n").font.size = Pt(11)
    meta.add_run("Confidencial — Uso Estratégico").font.size = Pt(10)
    for run in meta.runs:
        run.font.color.rgb = GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    # Score badges
    qa_score = qa.get("qa_score", 0)
    irr_alpha = irr.get("krippendorff_alpha", 0)

    badges = doc.add_paragraph()
    badges.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if qa_score:
        r_qa = badges.add_run(f"  QA Score: {qa_score}/100  ")
        r_qa.bold = True
        r_qa.font.size = Pt(11)
        r_qa.font.color.rgb = GREEN if qa_score >= 85 else ORANGE
    if irr_alpha:
        r_irr = badges.add_run(f"  IRR α: {irr_alpha:.2f}  ")
        r_irr.bold = True
        r_irr.font.size = Pt(11)
        r_irr.font.color.rgb = GREEN if irr_alpha >= 0.70 else RED

    doc.add_paragraph()
    prep = doc.add_paragraph()
    prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_prep = prep.add_run("Preparado por Sinergia Consulting Group · ARHIAX Dx v5.1")
    r_prep.font.size = Pt(9)
    r_prep.font.color.rgb = GRAY
    r_prep.italic = True

    doc.add_page_break()


# ── Table of Contents (static) ────────────────────────────────────────────────

def _build_toc(doc: Any):
    _heading(doc, "Contenido", 1, GREEN)
    sections = [
        ("1.",  "Resumen Ejecutivo"),
        ("2.",  "Contexto y Alcance"),
        ("3.",  "Instrumento Multi-Rater — Metodología y Resultados"),
        ("4.",  "Hallazgos Principales"),
        ("5.",  "Brechas de Percepción (Multi-Rater)"),
        ("6.",  "Cuellos de Botella Cuantificados"),
        ("7.",  "Análisis Bayesiano de Hipótesis"),
        ("8.",  "Recomendaciones Estratégicas"),
        ("9.",  "Roadmap de Implementación"),
        ("10.", "Próximos Pasos"),
        ("11.", "Control de Calidad y Gobernanza"),
        ("12.", "Metodología y Trazabilidad"),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for num, title in sections:
        row = tbl.add_row()
        row.cells[0].text = num
        row.cells[1].text = title
        row.cells[0].width = Cm(1.5)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_page_break()


# ── Section 1: Executive Summary ─────────────────────────────────────────────

def _build_executive_summary(doc: Any, redactor: dict, diagnostic: Any):
    _heading(doc, "1. Resumen Ejecutivo", 1, GREEN)

    summary = redactor.get("executive_summary", "")
    if not summary:
        summary = (
            f"El presente diagnóstico organizacional de {diagnostic.organization_name} "
            f"analiza el subproceso de {diagnostic.subprocess} en el sector {diagnostic.domain}. "
            f"El análisis fue realizado mediante el pipeline de 18 agentes IA de ARHIAX Dx v5.1, "
            f"con metodología Multi-Rater, análisis bayesiano y scoring psicométrico de 6 capas."
        )
    doc.add_paragraph(summary)

    context = redactor.get("context", "")
    if context:
        doc.add_paragraph()
        _heading(doc, "Contexto del Diagnóstico", 2)
        doc.add_paragraph(context)


# ── Section 2: Scope ─────────────────────────────────────────────────────────

def _build_scope(doc: Any, diagnostic: Any, g01: dict, g02: dict):
    _heading(doc, "2. Contexto y Alcance", 1, GREEN)

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    fields = [
        ("Organización",    diagnostic.organization_name),
        ("Sector",          diagnostic.domain),
        ("Subproceso",      diagnostic.subprocess),
        ("Tamaño",          f"{diagnostic.size_org} empleados" if diagnostic.size_org else "—"),
        ("Objetivo",        diagnostic.objective or "—"),
        ("Tipo diagnóstico",g01.get("diagnostic_type", "—")),
        ("Urgencia",        g01.get("urgency", "—")),
        ("Alcance",         g02.get("domain_config", {}).get("diagnostic_scope", "—")),
        ("Marcos de ref.",  ", ".join(g02.get("frameworks", [])[:4]) or "—"),
        ("Fecha",           datetime.now().strftime("%d/%m/%Y")),
    ]
    for key, val in fields:
        _add_kv_row(tbl, key, _safe_str(val))
    doc.add_paragraph()


# ── Section 3: Main Findings ──────────────────────────────────────────────────

def _build_findings(doc: Any, hallazgos: dict, redactor: dict):
    _heading(doc, "3. Hallazgos Principales", 1, GREEN)

    # Executive summary of findings
    exec_summary = hallazgos.get("executive_summary_findings", "")
    if exec_summary:
        p = doc.add_paragraph(exec_summary)
        p.runs[0].italic = True
        doc.add_paragraph()

    # Findings matrix table
    findings = hallazgos.get("findings_matrix", [])
    if not findings:
        # Fallback to redactor main_findings
        for f in redactor.get("main_findings", [])[:5]:
            if isinstance(f, dict):
                p = doc.add_paragraph(style="List Number")
                r = p.add_run(f.get("finding", str(f)))
                r.bold = True
                impact = f.get("impact", "")
                if impact:
                    doc.add_paragraph(f"   Impacto: {impact}")
            else:
                doc.add_paragraph(str(f), style="List Number")
        return

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    headers = ["ID", "Hallazgo", "Prioridad", "Confianza", "Impacto"]
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, "0A7F5A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

    priority_colors = {"CRITICA": "FEE2E2", "ALTA": "FEF3C7", "MEDIA": "DBEAFE", "BAJA": "F0FDF4"}
    for f in findings[:8]:
        row = tbl.add_row()
        row.cells[0].text = _safe_str(f.get("id"))
        row.cells[1].text = _safe_str(f.get("finding"))
        row.cells[2].text = _safe_str(f.get("priority"))
        conf = f.get("bayesian_confidence", 0)
        row.cells[3].text = f"{conf:.0%}" if conf else "—"
        row.cells[4].text = str(f.get("impact_score", "—"))
        prio = f.get("priority", "BAJA")
        bg = priority_colors.get(prio, "FFFFFF")
        _set_cell_bg(row.cells[2], bg)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # Problem statements
    pss = hallazgos.get("problem_statements", [])
    if pss:
        _heading(doc, "Declaraciones del Problema", 2)
        for ps in pss[:3]:
            stmt = ps.get("statement", str(ps)) if isinstance(ps, dict) else str(ps)
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stmt)


# ── Section 4: Perception Gaps (Multi-Rater) ─────────────────────────────────

def _build_perception_gaps(doc: Any, scoring: dict, bayesiano: dict, redactor: dict):
    _heading(doc, "4. Brechas de Percepción Multi-Rater", 1, GREEN)

    # Narrative
    gaps_text = redactor.get("perception_gaps", "")
    if gaps_text:
        doc.add_paragraph(gaps_text)
        doc.add_paragraph()

    # Role scores table
    role_scores = scoring.get("role_scores", {})
    if role_scores:
        _heading(doc, "Scores por Nivel Jerárquico", 2)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        for i, h in enumerate(["Rol", "Score", "Percepción", "Barra"]):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            _set_cell_bg(cell, "1A1A2E")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

        role_bg = {"Estratégico": "EEF2FF", "Táctico": "E0F2FE", "Operativo": "D1FAE5"}
        for role, data in role_scores.items():
            row = tbl.add_row()
            score = data.get("score", 0) if isinstance(data, dict) else 0
            perception = data.get("perception", "") if isinstance(data, dict) else ""
            row.cells[0].text = role
            row.cells[1].text = str(score)
            row.cells[2].text = perception.capitalize() if perception else "—"
            row.cells[3].text = _score_bar_text(int(score))
            _set_cell_bg(row.cells[0], role_bg.get(role, "FFFFFF"))
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()

    # Delta sigma
    delta = scoring.get("delta_sigma", {})
    max_gap = delta.get("max_gap", 0)
    if max_gap:
        _heading(doc, f"Brecha de Percepción Máxima (δσ = {max_gap:.2f})", 2)
        if max_gap > 2.0:
            p = doc.add_paragraph()
            r = p.add_run(f"⚠ BRECHA CRÍTICA DETECTADA (δσ = {max_gap:.2f} > 2.0)")
            r.bold = True
            r.font.color.rgb = RED
            doc.add_paragraph(
                "La diferencia de percepción entre niveles jerárquicos supera el umbral crítico. "
                "Este hallazgo fue escalado automáticamente a revisión humana (HIC MEDIUM)."
            )
        gap_pairs = delta.get("gap_pairs", [])
        if gap_pairs:
            tbl2 = doc.add_table(rows=1, cols=3)
            tbl2.style = "Table Grid"
            for i, h in enumerate(["Roles", "Delta σ", "Estado"]):
                cell = tbl2.rows[0].cells[i]
                cell.text = h
                _set_cell_bg(cell, "374151")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9)
            for gp in gap_pairs[:6]:
                row = tbl2.add_row()
                delta_val = gp.get("delta", 0)
                row.cells[0].text = _safe_str(gp.get("roles"))
                row.cells[1].text = f"{delta_val:.2f}"
                row.cells[2].text = "⚠ Crítico" if gp.get("critical") else "Normal"
                if gp.get("critical"):
                    _set_cell_bg(row.cells[2], "FEE2E2")
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(9)

    # Dimension scores
    dim_scores = scoring.get("dimension_scores", [])
    if dim_scores:
        doc.add_paragraph()
        _heading(doc, "Scores por Dimensión", 2)
        tbl3 = doc.add_table(rows=1, cols=4)
        tbl3.style = "Table Grid"
        for i, h in enumerate(["Dimensión", "Score", "Benchmark", "Brecha"]):
            cell = tbl3.rows[0].cells[i]
            cell.text = h
            _set_cell_bg(cell, "0A7F5A")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
        for d in dim_scores:
            row = tbl3.add_row()
            gap = d.get("gap", 0)
            row.cells[0].text = _safe_str(d.get("name", d.get("dimension")))
            row.cells[1].text = str(d.get("score", "—"))
            row.cells[2].text = str(d.get("benchmark", "—"))
            row.cells[3].text = f"{gap:+.0f}" if gap else "—"
            if gap and gap < -10:
                _set_cell_bg(row.cells[3], "FEE2E2")
            elif gap and gap > 0:
                _set_cell_bg(row.cells[3], "D1FAE5")
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)


# ── Section 5: Bottlenecks ────────────────────────────────────────────────────

def _build_bottlenecks(doc: Any, cuellos: dict, redactor: dict):
    _heading(doc, "5. Cuellos de Botella Cuantificados", 1, GREEN)

    summary = redactor.get("bottlenecks_summary", "")
    if summary:
        doc.add_paragraph(summary)
        doc.add_paragraph()

    total_loss = cuellos.get("total_opportunity_loss_usd_month", "")
    total_hours = cuellos.get("total_hours_lost_month", "")
    if total_loss or total_hours:
        p = doc.add_paragraph()
        if total_loss:
            r = p.add_run(f"Pérdida de oportunidad total: USD {total_loss}/mes  ")
            r.bold = True
            r.font.color.rgb = RED
        if total_hours:
            r2 = p.add_run(f"Horas perdidas: {total_hours} h/mes")
            r2.bold = True
            r2.font.color.rgb = ORANGE
        doc.add_paragraph()

    bottlenecks = cuellos.get("bottlenecks", [])
    if not bottlenecks:
        doc.add_paragraph("No se identificaron cuellos de botella cuantificados en este diagnóstico.")
        return

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    headers = ["ID", "Cuello de Botella", "Severidad", "Impacto", "Horas/mes", "USD/mes"]
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, "1A1A2E")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

    sev_colors = {"CRITICO": "FEE2E2", "ALTO": "FEF3C7", "MEDIO": "DBEAFE", "BAJO": "F0FDF4"}
    for b in bottlenecks[:7]:
        row = tbl.add_row()
        sev = _safe_str(b.get("severity", "MEDIO"))
        row.cells[0].text = _safe_str(b.get("id"))
        row.cells[1].text = _safe_str(b.get("name"))
        row.cells[2].text = sev
        row.cells[3].text = str(b.get("impact_score", "—"))
        row.cells[4].text = str(b.get("estimated_hours_lost_month", "—"))
        row.cells[5].text = str(b.get("estimated_cost_usd_month", "—"))
        _set_cell_bg(row.cells[2], sev_colors.get(sev, "FFFFFF"))
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    # Quick wins
    quick_wins = cuellos.get("quick_fix_opportunities", [])
    if quick_wins:
        doc.add_paragraph()
        _heading(doc, "Oportunidades de Mejora Rápida", 2)
        for qw in quick_wins[:4]:
            if isinstance(qw, dict):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{qw.get('bottleneck', '')}: ").bold = True
                p.add_run(f"{qw.get('fix', '')} — Esfuerzo: {qw.get('effort', '')} · Impacto: {qw.get('impact', '')}")
            else:
                doc.add_paragraph(str(qw), style="List Bullet")


# ── Section 3: Multi-Rater Instrument ────────────────────────────────────────

def _build_instrument(doc: Any, g09a: dict, g09c: dict, irr: dict, psico: dict, scoring: dict):
    """
    Section 3 — Instrumento Multi-Rater.
    Shows: dimensions mapped to hypotheses, questions per dimension,
    respondents by role, IRR, Cronbach alpha, and per-dimension score bars.
    This section justifies WHY those questions were asked and WHAT they measured.
    """
    _heading(doc, "3. Instrumento Multi-Rater — Metodología y Resultados", 1, GREEN)

    # ── Intro paragraph ───────────────────────────────────────────────────────
    instrument_name = g09a.get("instrument_name", "Instrumento Multi-Rater")
    methodology = g09a.get("methodology", {})
    standard = methodology.get("standard", "Kirkpatrick (1994) adaptado para diagnóstico organizacional")
    design_principle = methodology.get("design_principle", "")

    intro = doc.add_paragraph()
    intro.add_run(f"{instrument_name}. ").bold = True
    intro.add_run(
        f"El instrumento fue diseñado específicamente para este diagnóstico siguiendo el estándar "
        f"{standard}. Cada dimensión mapea directamente a una hipótesis diagnóstica identificada "
        f"en la fase de análisis de brechas, garantizando que cada pregunta tiene un propósito "
        f"verificable y una señal esperada por nivel jerárquico."
    )
    if design_principle:
        p2 = doc.add_paragraph()
        p2.add_run("Principio de diseño: ").bold = True
        p2.add_run(design_principle)

    doc.add_paragraph()

    # ── Reliability metrics ───────────────────────────────────────────────────
    _heading(doc, "Métricas de Confiabilidad del Instrumento", 2)

    irr_alpha   = irr.get("krippendorff_alpha", 0)
    irr_status  = irr.get("irr_status", "")
    cronbach    = psico.get("cronbach_alpha_overall", 0)
    consistency = psico.get("internal_consistency", "")
    reliability = psico.get("instrument_reliability", "")

    tbl_rel = doc.add_table(rows=0, cols=3)
    tbl_rel.style = "Table Grid"
    # Header
    hdr = tbl_rel.add_row()
    for i, h in enumerate(["Métrica", "Valor", "Interpretación"]):
        hdr.cells[i].text = h
        _set_cell_bg(hdr.cells[i], "1A1A2E")
        for run in hdr.cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

    metrics = []
    if irr_alpha:
        ok = irr_alpha >= 0.70
        metrics.append((
            "α Krippendorff (IRR)",
            f"{irr_alpha:.3f}",
            f"{irr_status} — Acuerdo inter-evaluador {'aceptable' if ok else 'bajo'} "
            f"(umbral mínimo: 0.70, Krippendorff 2004)",
            "D1FAE5" if ok else "FEE2E2",
        ))
    if cronbach:
        ok = cronbach >= 0.70
        metrics.append((
            "α Cronbach (consistencia interna)",
            f"{cronbach:.3f}",
            f"{consistency} — Fiabilidad {reliability.lower() if reliability else ''} "
            f"(umbral mínimo: 0.70, Nunnally 1978)",
            "D1FAE5" if ok else "FEE2E2",
        ))

    # Per-dimension Cronbach
    cronbach_by_dim = psico.get("cronbach_by_dimension", {})
    irr_by_dim = irr.get("by_dimension", {})

    for label, val, interp, bg in metrics:
        row = tbl_rel.add_row()
        row.cells[0].text = label
        row.cells[1].text = val
        row.cells[2].text = interp
        _set_cell_bg(row.cells[1], bg)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # ── Respondents by role ───────────────────────────────────────────────────
    role_scores = scoring.get("role_scores", {})
    if role_scores:
        _heading(doc, "Participación por Nivel Jerárquico", 2)
        tbl_roles = doc.add_table(rows=0, cols=4)
        tbl_roles.style = "Table Grid"
        hdr2 = tbl_roles.add_row()
        for i, h in enumerate(["Rol", "Respondentes", "Score Promedio", "Percepción"]):
            hdr2.cells[i].text = h
            _set_cell_bg(hdr2.cells[i], "0A7F5A")
            for run in hdr2.cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

        role_bg = {"Estratégico": "EEF2FF", "Táctico": "E0F2FE", "Operativo": "D1FAE5"}
        for role, data in role_scores.items():
            if not isinstance(data, dict):
                continue
            row = tbl_roles.add_row()
            score = data.get("score", 0)
            n = data.get("n_responses", "—")
            perception = data.get("perception", "—")
            row.cells[0].text = role
            row.cells[1].text = str(n)
            row.cells[2].text = f"{score}/100  {_score_bar_text(int(score))}"
            row.cells[3].text = perception.capitalize() if perception else "—"
            _set_cell_bg(row.cells[0], role_bg.get(role, "FFFFFF"))
            score_color = "D1FAE5" if score >= 70 else ("FEF3C7" if score >= 50 else "FEE2E2")
            _set_cell_bg(row.cells[2], score_color)
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()

    # ── Dimensions table — the core of this section ───────────────────────────
    _heading(doc, "Dimensiones del Instrumento y Resultados", 2)

    p_dim_intro = doc.add_paragraph()
    p_dim_intro.add_run(
        "Cada dimensión fue diseñada para verificar una hipótesis específica identificada en el "
        "análisis de brechas. La columna 'Señal esperada' describe qué patrón de respuestas "
        "confirmaría la hipótesis. La columna 'Score observado' muestra el resultado real."
    )
    p_dim_intro.runs[0].font.size = Pt(9)
    p_dim_intro.runs[0].font.color.rgb = GRAY
    doc.add_paragraph()

    dimensions = g09a.get("dimensions", [])
    dim_scores_list = scoring.get("dimension_scores", [])
    dim_scores_map = {d.get("dimension", ""): d for d in dim_scores_list}

    # Also build from g09c dimension_coverage
    dim_coverage = g09c.get("dimension_coverage", {})

    for dim in dimensions:
        dim_id   = dim.get("id", "")
        dim_name = dim.get("name", dim_id)
        hyp_id   = dim.get("hypothesis_mapped", "")
        hyp_text = dim.get("hypothesis_text", "")
        expected_true  = dim.get("expected_pattern_if_true", "")
        expected_false = dim.get("expected_pattern_if_false", "")
        rationale = dim.get("rationale", "")

        # Score data
        score_data = dim_scores_map.get(dim_id, {})
        dim_score  = score_data.get("score")
        benchmark  = score_data.get("benchmark", 75)
        gap        = score_data.get("gap")

        # Coverage data
        coverage = dim_coverage.get(dim_id, {}) if isinstance(dim_coverage, dict) else {}
        n_likert  = coverage.get("questions_count", 0)
        n_open    = coverage.get("open_questions", 0)
        n_reverse = coverage.get("reverse_scored", 0)

        # Cronbach for this dimension
        alpha_dim = cronbach_by_dim.get(dim_id, "—")
        irr_dim   = irr_by_dim.get(dim_id, {})
        irr_dim_val = irr_dim.get("alpha", "—") if isinstance(irr_dim, dict) else "—"

        # Dimension header
        p_dh = doc.add_paragraph()
        r_dh = p_dh.add_run(f"  {dim_id}: {dim_name}  ")
        r_dh.bold = True
        r_dh.font.size = Pt(11)
        r_dh.font.color.rgb = DARK
        if hyp_id:
            r_hyp = p_dh.add_run(f"  → {hyp_id}")
            r_hyp.bold = True
            r_hyp.font.size = Pt(10)
            r_hyp.font.color.rgb = GREEN

        # Dimension detail table
        tbl_dim = doc.add_table(rows=0, cols=2)
        tbl_dim.style = "Table Grid"

        rows_data = []
        if hyp_text:
            rows_data.append(("Hipótesis que verifica", hyp_text))
        if rationale:
            rows_data.append(("Justificación de la dimensión", rationale))
        if expected_true:
            rows_data.append(("Señal esperada si hipótesis es verdadera", expected_true))
        if expected_false:
            rows_data.append(("Señal si hipótesis es falsa", expected_false))

        struct = f"{n_likert} ítems Likert 1-5"
        if n_open:
            struct += f" + {n_open} pregunta{'s' if n_open > 1 else ''} abierta{'s' if n_open > 1 else ''}"
        if n_reverse:
            struct += f" · {n_reverse} reverse-scored (control sesgo aquiescencia)"
        if struct:
            rows_data.append(("Estructura del instrumento", struct))

        reliability_str = ""
        if alpha_dim != "—":
            reliability_str += f"α Cronbach = {alpha_dim}"
        if irr_dim_val != "—":
            reliability_str += f"  ·  α Krippendorff = {irr_dim_val}"
        if reliability_str:
            rows_data.append(("Confiabilidad de la dimensión", reliability_str))

        if dim_score is not None:
            score_str = f"{dim_score}/100"
            if benchmark:
                score_str += f"  (benchmark sector: {benchmark}/100)"
            if gap is not None:
                score_str += f"  ·  Brecha: {gap:+.0f} puntos"
            score_str += f"\n{_score_bar_text(int(dim_score))}"
            rows_data.append(("Score observado (corregido)", score_str))

        for key, val in rows_data:
            _add_kv_row(tbl_dim, key, val)

        doc.add_paragraph()

    # ── Questions summary table ───────────────────────────────────────────────
    _heading(doc, "Resumen de Preguntas por Dimensión", 2)

    questions = g09a.get("questions", [])
    if questions:
        tbl_q = doc.add_table(rows=1, cols=5)
        tbl_q.style = "Table Grid"
        for i, h in enumerate(["ID", "Dimensión", "Hipótesis", "Tipo", "Roles"]):
            cell = tbl_q.rows[0].cells[i]
            cell.text = h
            _set_cell_bg(cell, "374151")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(8)

        for q in questions:
            row = tbl_q.add_row()
            qid   = q.get("id", "")
            qtype = q.get("type", "")
            roles = ", ".join(r[:3] for r in q.get("roles", []))
            hyp   = q.get("hypothesis_tested", "")
            rev   = " ↔" if q.get("reverse_scored") else ""

            row.cells[0].text = qid
            row.cells[1].text = q.get("dimension", "")
            row.cells[2].text = hyp
            row.cells[3].text = ("Likert 1-5" + rev) if "likert" in qtype else "Abierta"
            row.cells[4].text = roles

            if q.get("reverse_scored"):
                _set_cell_bg(row.cells[3], "FEF3C7")
            elif "open" in qtype:
                _set_cell_bg(row.cells[3], "DBEAFE")

            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)

        # Legend
        doc.add_paragraph()
        legend = doc.add_paragraph()
        legend.add_run("↔ Reverse-scored: ").bold = True
        legend.add_run(
            "el score se invierte (6 − respuesta) antes de calcular el promedio. "
            "Controla el sesgo de aquiescencia (Paulhus, 1991). "
            "Score alto = situación positiva en todas las preguntas."
        )
        legend.runs[-1].font.size = Pt(8)
        legend.runs[-1].font.color.rgb = GRAY


# ── Section 12: Methodology & Traceability ───────────────────────────────────

def _build_methodology(doc: Any, g09a: dict, g09c: dict, g05: dict, irr: dict, psico: dict):
    """
    Section 12 — Metodología y Trazabilidad.
    Full justification of the instrument design, scoring method, and audit trail.
    This section answers: why these questions, why this number, why this scoring.
    """
    _heading(doc, "12. Metodología y Trazabilidad", 1, GREEN)

    doc.add_paragraph(
        "Esta sección documenta la justificación metodológica completa del diagnóstico. "
        "Cualquier hallazgo, score o recomendación en este informe puede ser trazado hasta "
        "las respuestas individuales de la encuesta, las hipótesis originales y las reglas "
        "de gobernanza que gobernaron la ejecución del pipeline."
    )
    doc.add_paragraph()

    # ── 12.1 Instrument design rationale ─────────────────────────────────────
    _heading(doc, "12.1 Diseño del Instrumento", 2)

    methodology = g09a.get("methodology", {})
    standard    = methodology.get("standard", "Kirkpatrick (1994) adaptado")
    irr_target  = methodology.get("irr_target", "α Krippendorff ≥ 0.70 (Krippendorff, 2004)")
    rev_scoring = methodology.get("reverse_scoring", "≥1 ítem reverse-scored por dimensión (Paulhus, 1991)")
    role_diff   = methodology.get("role_differentiation", "")
    design_p    = methodology.get("design_principle", "")

    tbl_m = doc.add_table(rows=0, cols=2)
    tbl_m.style = "Table Grid"
    rows_m = [
        ("Estándar metodológico",       standard),
        ("Principio de diseño",         design_p or "Verificabilidad por rol: cada dimensión produce señal diferenciada si la hipótesis es verdadera"),
        ("Objetivo IRR",                irr_target),
        ("Control de sesgo",            rev_scoring),
        ("Diferenciación de roles",     role_diff or "Preguntas asignadas por rol según nivel de acceso a la información del proceso"),
        ("Estructura del instrumento",  "15 ítems Likert 1-5 + 3 preguntas abiertas = 18 ítems totales"),
        ("Dimensiones",                 "4 dimensiones, una por hipótesis (DIM-01→H01, DIM-02→H02, DIM-03→H03, DIM-04→H04)"),
        ("Mínimo ítems/dimensión",      "3 ítems Likert (mínimo para α Cronbach confiable, Nunnally 1978)"),
    ]
    for key, val in rows_m:
        if val:
            _add_kv_row(tbl_m, key, val)
    doc.add_paragraph()

    # ── 12.2 Hypothesis traceability ─────────────────────────────────────────
    _heading(doc, "12.2 Trazabilidad Hipótesis → Dimensión → Preguntas", 2)

    doc.add_paragraph(
        "Cada hipótesis generada en la fase de análisis de brechas (G05) fue mapeada a una "
        "dimensión del instrumento. Las preguntas de esa dimensión fueron diseñadas para "
        "producir una señal diferenciada por rol si la hipótesis es verdadera. "
        "Esta cadena es verificable: G05 → G09a → G10a → G11a."
    )
    doc.add_paragraph()

    hypotheses = g05.get("hypotheses", [])
    dimensions = g09a.get("dimensions", [])
    dim_map    = {d.get("hypothesis_mapped", ""): d for d in dimensions}

    if hypotheses:
        tbl_h = doc.add_table(rows=1, cols=4)
        tbl_h.style = "Table Grid"
        for i, h in enumerate(["Hipótesis", "Prior P(H)", "Dimensión", "Condición de falsificación"]):
            cell = tbl_h.rows[0].cells[i]
            cell.text = h
            _set_cell_bg(cell, "0A7F5A")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

        for hyp in hypotheses:
            hid   = hyp.get("id", "")
            htext = hyp.get("hypothesis", "")
            prior = hyp.get("prior_probability", 0)
            dim   = hyp.get("dimension_to_measure", dim_map.get(hid, {}).get("id", ""))
            ev    = hyp.get("evidence_needed", "")

            # Get falsification from expected_signals
            signals = hyp.get("expected_signals", {})
            falsif  = signals.get("falsification_condition", ev)

            row = tbl_h.add_row()
            row.cells[0].text = f"{hid}: {htext[:70]}{'…' if len(htext) > 70 else ''}"
            row.cells[1].text = f"{prior:.2f}"
            row.cells[2].text = dim
            row.cells[3].text = falsif[:80] if falsif else "—"
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
        doc.add_paragraph()

    # ── 12.3 Scoring method ───────────────────────────────────────────────────
    _heading(doc, "12.3 Método de Scoring", 2)

    tbl_s = doc.add_table(rows=0, cols=2)
    tbl_s.style = "Table Grid"
    scoring_rows = [
        ("Normalización Likert",        "score = (valor − 1) / 4 × 100  →  1=0, 2=25, 3=50, 4=75, 5=100"),
        ("Corrección reverse-scoring",  "score_corregido = 6 − valor_original  (aplicado antes de normalizar)"),
        ("Aplicado por",                "pipeline_runner.py — código Python determinista, no delegado al LLM"),
        ("Score por dimensión",         "Promedio de scores corregidos de todos los ítems de esa dimensión"),
        ("Score por rol",               "Promedio de scores de todas las dimensiones para ese rol"),
        ("Score global",                "Promedio ponderado de dimensiones (peso igual: 0.25 cada una)"),
        ("Delta sigma (δσ)",            "|score_rol_A − score_rol_B| / 20  ·  δσ > 2.0 = brecha crítica"),
        ("Benchmark sectorial",         "Score de referencia del sector (70-78 para sector maduro)"),
        ("Percentil",                   "Posición del score global vs distribución del sector"),
    ]
    for key, val in scoring_rows:
        _add_kv_row(tbl_s, key, val)
    doc.add_paragraph()

    # ── 12.4 Audit trail ─────────────────────────────────────────────────────
    _heading(doc, "12.4 Cadena de Auditoría", 2)

    audit_cert = g09c.get("audit_certificate", {})
    methodology_compliance = g09c.get("methodology_compliance", {})

    doc.add_paragraph(
        "El diagnóstico fue ejecutado bajo el framework de gobernanza ARHIAX Dx v5.1. "
        "Cada decisión del pipeline está registrada en el ledger de evidencia append-only "
        "con hashes SHA-256 encadenados y firma criptográfica Ed25519."
    )
    doc.add_paragraph()

    tbl_a = doc.add_table(rows=0, cols=2)
    tbl_a.style = "Table Grid"
    audit_rows = [
        ("Validado por",                audit_cert.get("validated_by", "G09c — ARHIAX Dx")),
        ("Estándar de validación",      audit_cert.get("methodology", "Kirkpatrick (1994) adaptado")),
        ("Estándar IRR",                audit_cert.get("irr_standard", "Krippendorff (2004) α ≥ 0.70")),
        ("Trazabilidad H→Q",            audit_cert.get("hypothesis_traceability", "H01-H04 → DIM-01-DIM-04 → Q01-Q15 + QA01-QA03")),
        ("Hipótesis trazables",         "Sí" if methodology_compliance.get("hypothesis_traceability") else "—"),
        ("Reverse-scoring documentado", "Sí" if methodology_compliance.get("reverse_scoring_present") else "—"),
        ("Rationale documentado",       "Sí" if methodology_compliance.get("rationale_documented") else "—"),
        ("Firma criptográfica",         "Ed25519 — verificable con clave pública del agente"),
        ("Ledger de evidencia",         "Append-only, hashes SHA-256 encadenados"),
        ("Anonimización respondentes",  "Hash SHA-256 por respondente — datos personales nunca almacenados"),
        ("Retención de datos",          "30 días — AUTO_DELETE posterior (política de minimización)"),
        ("Pipeline",                    "ARHIAX Dx v5.1 — 18 agentes IA — Gemini 2.5 Flash/Pro"),
    ]
    for key, val in audit_rows:
        if val and val != "—":
            _add_kv_row(tbl_a, key, val)

    doc.add_paragraph()

    # ── 12.5 References ───────────────────────────────────────────────────────
    _heading(doc, "12.5 Referencias Metodológicas", 2)

    refs = [
        "Kirkpatrick, D.L. (1994). Evaluating Training Programs. Berrett-Koehler.",
        "Nunnally, J.C. (1978). Psychometric Theory (2nd ed.). McGraw-Hill.",
        "Krippendorff, K. (2004). Content Analysis: An Introduction to Its Methodology. Sage.",
        "Paulhus, D.L. (1991). Measurement and control of response bias. In J.P. Robinson et al. (Eds.), Measures of Personality and Social Psychological Attitudes. Academic Press.",
        "Bayes, T. (1763). An Essay towards solving a Problem in the Doctrine of Chances. Philosophical Transactions of the Royal Society.",
        "Sinergia Consulting Group (2026). ARHIAX Dx Briefing v5.1 — Sistema Agentico de Diagnóstico Organizacional Gobernado.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(ref).font.size = Pt(9)

def _build_bayesian(doc: Any, bayesiano: dict, g05: dict):
    _heading(doc, "7. Análisis Bayesiano de Hipótesis", 1, GREEN)

    summary = bayesiano.get("bayesian_summary", "")
    if summary:
        doc.add_paragraph(summary)
        doc.add_paragraph()

    analysis = bayesiano.get("bayesian_analysis", [])
    if analysis:
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        for i, h in enumerate(["ID", "Hipótesis", "Prior P(H)", "Posterior P(H|E)", "Estado", "Evidencia clave"]):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            _set_cell_bg(cell, "0A7F5A")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(8)

        for item in analysis[:8]:
            row = tbl.add_row()
            confirmed = item.get("confirmed", False)
            posterior = item.get("posterior_probability", 0)
            hid       = item.get("hypothesis_id", item.get("id", ""))
            hyp_text  = _safe_str(item.get("hypothesis"))
            ev_score  = item.get("evidence_score", {})
            ev_list   = item.get("evidence_used", [])
            ev_summary = ""
            if ev_score and isinstance(ev_score, dict):
                ev_summary = ev_score.get("reasoning", "")
            elif ev_list:
                ev_summary = ev_list[0] if ev_list else ""

            row.cells[0].text = hid
            row.cells[1].text = hyp_text[:60] + ("…" if len(hyp_text) > 60 else "")
            row.cells[2].text = f"{item.get('prior_probability', 0):.2f}"
            row.cells[3].text = f"{posterior:.2f}"
            row.cells[4].text = "✓ Confirmada" if confirmed else "✗ Rechazada"
            row.cells[5].text = ev_summary[:60] if ev_summary else "—"
            _set_cell_bg(row.cells[4], "D1FAE5" if confirmed else "FEE2E2")
            if posterior >= 0.85:
                _set_cell_bg(row.cells[3], "D1FAE5")
            elif posterior < 0.50:
                _set_cell_bg(row.cells[3], "FEE2E2")
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
        doc.add_paragraph()

    # Causality map
    causality = bayesiano.get("causality_map", {})
    root_causes = causality.get("root_causes", [])
    if root_causes:
        _heading(doc, "Mapa de Causalidad", 2)
        p_rc = doc.add_paragraph()
        p_rc.add_run("Causas raíz identificadas: ").bold = True
        p_rc.add_run(", ".join(root_causes))
        for rel in causality.get("relationships", [])[:4]:
            if isinstance(rel, dict):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{rel.get('cause', '')} → {rel.get('effect', '')} ").bold = True
                p.add_run(f"(fuerza: {rel.get('strength', 0):.2f}): {rel.get('explanation', '')}")
                p.runs[-1].font.size = Pt(9)
        doc.add_paragraph()

    # Critical perception gaps from bayesian
    gaps = bayesiano.get("critical_perception_gaps", [])
    if gaps:
        _heading(doc, "Brechas de Percepción Críticas", 2)
        for gap in gaps:
            p = doc.add_paragraph(style="List Bullet")
            dim   = gap.get("dimension", "")
            delta = gap.get("delta_sigma", 0)
            roles = gap.get("roles", "")
            interp = gap.get("interpretation", "")
            p.add_run(f"{dim} — {roles} (δσ = {delta:.1f}): ").bold = True
            p.add_run(interp)
            if gap.get("escalate"):
                r = p.add_run(" [ESCALADO HIC MEDIUM]")
                r.font.color.rgb = RED
                r.bold = True


# ── Section 8: Recommendations ───────────────────────────────────────────────

def _build_recommendations(doc: Any, hallazgos: dict, redactor: dict):
    _heading(doc, "8. Recomendaciones Estratégicas", 1, GREEN)

    recs = hallazgos.get("strategic_recommendations") or redactor.get("strategic_recommendations", [])
    if not recs:
        doc.add_paragraph("No se generaron recomendaciones en este diagnóstico.")
        return

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    for i, h in enumerate(["#", "Recomendación", "Plazo", "Impacto", "Inversión"]):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, "0A7F5A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

    timeframe_colors = {"90_dias": "D1FAE5", "180_dias": "DBEAFE", "365_dias": "FEF3C7"}
    for rec in recs[:6]:
        if not isinstance(rec, dict):
            row = tbl.add_row()
            row.cells[1].text = str(rec)
            continue
        row = tbl.add_row()
        tf = _safe_str(rec.get("timeframe", rec.get("expected_roi", "")))
        row.cells[0].text = str(rec.get("priority", rec.get("id", "—")))
        rec_text = _safe_str(rec.get("recommendation"))
        row.cells[1].text = rec_text[:100] + ("…" if len(rec_text) > 100 else "")
        row.cells[2].text = tf.replace("_", " ")
        row.cells[3].text = _safe_str(rec.get("expected_impact", rec.get("expected_roi", "—")))[:60]
        row.cells[4].text = _safe_str(rec.get("investment_level", "—"))
        bg = timeframe_colors.get(tf, "FFFFFF")
        _set_cell_bg(row.cells[2], bg)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)


# ── Section 8: Roadmap ────────────────────────────────────────────────────────

def _build_roadmap(doc: Any, redactor: dict, optimizador: dict):
    _heading(doc, "8. Roadmap de Implementación", 1, GREEN)

    roadmap = redactor.get("roadmap") or optimizador.get("roadmap", {})
    if not roadmap:
        doc.add_paragraph("Roadmap no disponible en este diagnóstico.")
        return

    periods = [
        ("days_90",  "Primeros 90 días",  "Estabilización",  "D1FAE5", "0A7F5A"),
        ("days_180", "90 a 180 días",     "Optimización",    "DBEAFE", "2563EB"),
        ("days_365", "180 a 365 días",    "Transformación",  "FEF3C7", "D97706"),
    ]

    for key, label, default_theme, bg, header_color in periods:
        period = roadmap.get(key, {})
        if not period:
            continue

        # Period header
        p_header = doc.add_paragraph()
        r = p_header.add_run(f"  {label}  ")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(
            int(header_color[0:2], 16),
            int(header_color[2:4], 16),
            int(header_color[4:6], 16),
        )

        theme = ""
        actions = []
        outcome = ""

        if isinstance(period, dict):
            theme = period.get("theme", default_theme)
            actions = period.get("actions", [])
            outcome = period.get("expected_outcome", "")
        elif isinstance(period, list):
            actions = period

        if theme:
            p_theme = doc.add_paragraph()
            p_theme.add_run(f"Tema: {theme}").italic = True

        for action in actions[:5]:
            doc.add_paragraph(str(action), style="List Bullet")

        if outcome:
            p_out = doc.add_paragraph()
            p_out.add_run("Resultado esperado: ").bold = True
            p_out.add_run(outcome)

        doc.add_paragraph()


# ── Section 9: Next Steps ─────────────────────────────────────────────────────

def _build_next_steps(doc: Any, redactor: dict):
    _heading(doc, "9. Próximos Pasos Inmediatos", 1, GREEN)

    steps = redactor.get("next_steps", [])
    if not steps:
        doc.add_paragraph("Definir próximos pasos con el equipo directivo.")
        return

    for i, step in enumerate(steps[:6], 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(str(step))


# ── Section 10: QA & Governance ──────────────────────────────────────────────

def _build_qa_governance(doc: Any, qa: dict, irr: dict, psico: dict, diagnostic: Any):
    _heading(doc, "10. Control de Calidad y Gobernanza", 1, GREEN)

    qa_score = qa.get("qa_score", 0)
    if qa_score:
        p = doc.add_paragraph()
        r = p.add_run(f"Score QA: {qa_score}/100  ")
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = GREEN if qa_score >= 85 else ORANGE
        status = "APROBADO" if qa_score >= 85 else "REQUIERE REVISIÓN"
        r2 = p.add_run(f"[{status}]")
        r2.bold = True
        r2.font.color.rgb = GREEN if qa_score >= 85 else RED

    # QA dimensions table
    dims = qa.get("quality_dimensions", {})
    if dims:
        doc.add_paragraph()
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        for i, h in enumerate(["Dimensión QA", "Score", "Barra"]):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            _set_cell_bg(cell, "374151")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
        for dim_name, dim_data in dims.items():
            row = tbl.add_row()
            score = dim_data.get("score", 0) if isinstance(dim_data, dict) else 0
            max_s = dim_data.get("max", 20) if isinstance(dim_data, dict) else 20
            row.cells[0].text = dim_name.replace("_", " ").title()
            row.cells[1].text = f"{score}/{max_s}"
            row.cells[2].text = _score_bar_text(score, max_s)
            if score >= max_s * 0.85:
                _set_cell_bg(row.cells[1], "D1FAE5")
            elif score < max_s * 0.70:
                _set_cell_bg(row.cells[1], "FEE2E2")
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()

    # IRR & Psychometrics
    irr_alpha = irr.get("krippendorff_alpha", 0)
    cronbach = psico.get("cronbach_alpha_overall", 0)
    if irr_alpha or cronbach:
        _heading(doc, "Métricas Psicométricas", 2)
        tbl2 = doc.add_table(rows=0, cols=2)
        tbl2.style = "Table Grid"
        if irr_alpha:
            _add_kv_row(tbl2, "Krippendorff Alpha (IRR)", f"{irr_alpha:.3f} — {irr.get('irr_status', '')}")
        if cronbach:
            _add_kv_row(tbl2, "Alpha de Cronbach", f"{cronbach:.3f} — {psico.get('internal_consistency', '')}")
        doc.add_paragraph()

    # Governance statement
    doc.add_paragraph()
    gov_p = doc.add_paragraph()
    gov_p.add_run("Declaración de Gobernanza: ").bold = True
    gov_p.add_run(
        "Este informe fue generado por ARHIAX Dx v5.1 bajo el framework de gobernanza de Sinergia Consulting Group. "
        "El pipeline de 18 agentes IA fue ejecutado con nivel de autonomía A1/A2, con evaluación de 18 reglas de "
        "gobernanza, firma criptográfica Ed25519 y ledger de evidencia append-only. "
        "Todos los datos de respondentes fueron procesados de forma anónima."
    )
    gov_p.runs[-1].font.size = Pt(9)
    gov_p.runs[-1].font.color.rgb = GRAY


# ── Main entry point ──────────────────────────────────────────────────────────

def build_docx(diagnostic: Any, stages: list[Any]) -> bytes:
    """
    Build the executive Word report from pipeline outputs.
    Returns raw bytes of the .docx file.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    # Extract outputs from completed stages
    outputs: dict[str, dict] = {}
    for stage in stages:
        if not stage.output or not isinstance(stage.output, dict):
            continue
        # stage.output structure: {"tool": "...", "model_used": "...", "output": {...}}
        # Extract the nested "output" dict, fallback to stage.output if no nesting
        tool_output = stage.output.get("output", {})
        if tool_output and isinstance(tool_output, dict):
            outputs[stage.tool_name] = tool_output
        elif stage.output and "tool" not in stage.output:
            # Legacy format: stage.output is the tool output directly
            outputs[stage.tool_name] = stage.output

    # Pull each agent output (with fallbacks)
    g01       = outputs.get("g01_receptor", {})
    g02       = outputs.get("g02_configurador", {})
    cuellos   = outputs.get("g07_cuellos", {})
    optimizador = outputs.get("g08_optimizador", {})
    scoring   = outputs.get("g10a_scoring", {})
    psico     = outputs.get("g10b_psicometria", {})
    bayesiano = outputs.get("g11a_bayesiano", {})
    hallazgos = outputs.get("g12_hallazgos", {})
    redactor  = outputs.get("g13_redactor", {})
    qa        = outputs.get("g14_qa_control", {})
    irr       = outputs.get("irr_calculator", {})

    doc = Document()

    # ── Global styles ─────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = GREEN

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = DARK

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Build sections ────────────────────────────────────────────────────────
    _build_cover(doc, diagnostic, qa, irr)
    _build_toc(doc)
    _build_executive_summary(doc, redactor, diagnostic)
    doc.add_page_break()
    _build_scope(doc, diagnostic, g01, g02)
    doc.add_page_break()
    _build_findings(doc, hallazgos, redactor)
    doc.add_page_break()
    _build_perception_gaps(doc, scoring, bayesiano, redactor)
    doc.add_page_break()
    _build_bottlenecks(doc, cuellos, redactor)
    doc.add_page_break()
    _build_bayesian(doc, bayesiano)
    doc.add_page_break()
    _build_recommendations(doc, hallazgos, redactor)
    doc.add_page_break()
    _build_roadmap(doc, redactor, optimizador)
    doc.add_page_break()
    _build_next_steps(doc, redactor)
    doc.add_page_break()
    _build_qa_governance(doc, qa, irr, psico, diagnostic)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run(
        "ARHIAX Dx v5.1  ·  Sinergia Consulting Group  ·  Confidencial — Uso Estratégico"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY
    r.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
